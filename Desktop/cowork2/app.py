import inspect
from datetime import datetime, timedelta, time
from pathlib import Path

import pandas as pd
import streamlit as st
from utils.github_config import configure_github_env

# ── 背景排程器（每 60 秒檢查一次到期排程） ──────────────────────────────
# 必須在其他 import 之前啟動，且使用 st.cache_resource 確保只啟動一次
from utils.scheduler_daemon import start_background_scheduler
start_background_scheduler()
# ─────────────────────────────────────────────────────────────────────────

from utils.loaders import (
    load_sources,
    save_sources,
    load_experts,
    save_experts,
    load_profiles,
    save_profiles,
    load_formats,
    save_formats,
    load_insights,
    save_insights,
    load_auto_export,
    save_auto_export,
    get_source_categories,
    get_expert_categories,
    source_to_editor_row,
    editor_row_to_source,
    expert_to_editor_row,
    editor_row_to_expert,
    display_expert_name,
    load_category_keywords,
    save_category_keywords,
    DEFAULT_CATEGORY_KEYWORDS,
)
import os
import google.generativeai as genai

# Configure Gemini API before importing report_engine
if "GOOGLE_API_KEY" in st.secrets:
    os.environ["GOOGLE_API_KEY"] = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

github_ready, github_source = configure_github_env()

import report_engine

try:
    from utils import google_drive as google_drive_utils
except Exception:
    google_drive_utils = None


APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

st.set_page_config(page_title="Briefings", layout="wide")
st.title("Briefings")
st.caption("本機版 AI 情報簡報系統")
if github_ready:
    st.sidebar.success(f"GitHub 已接上（{github_source}）")
else:
    st.sidebar.info("GitHub 尚未設定")


# =========================================================
# Helpers
# =========================================================
def _now_str():
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _clean_batch_df(df: pd.DataFrame):
    if df is None or df.empty:
        return []
    rows = []
    for _, row in df.iterrows():
        item = {}
        has_value = False
        for col in df.columns:
            val = row[col]
            if pd.isna(val):
                val = ""
            if isinstance(val, str):
                val = val.strip()
            if val not in ["", None]:
                has_value = True
            item[col] = val
        if has_value:
            rows.append(item)
    return rows


def _profiles_map(profiles):
    result = {}
    for p in profiles:
        if isinstance(p, dict):
            name = p.get("name") or p.get("id") or p.get("title") or "Default"
            result[name] = p
        else:
            result[str(p)] = p
    return result


def _filter_sources_by_selection(sources, selected_categories, selected_names):
    filtered = []
    for s in sources:
        if not s.get("enabled", True):
            continue
        cats = s.get("category", []) or []
        if selected_categories and not any(c in selected_categories for c in cats):
            continue
        if selected_names and s.get("name") not in selected_names:
            continue
        filtered.append(s)
    return filtered


def _filter_experts_by_selection(experts, selected_categories, selected_names):
    filtered = []
    for e in experts:
        if not e.get("enabled", True):
            continue
        cats = e.get("category", []) or []
        name = display_expert_name(e)
        if selected_categories and not any(c in selected_categories for c in cats):
            continue
        if selected_names and name not in selected_names:
            continue
        filtered.append(e)
    return filtered


def _build_source_fetch_preview_rows(sources):
    category_keywords = load_category_keywords()
    rows = []

    for src in sources:
        cats = src.get("category", []) or []
        if isinstance(cats, str):
            cats = [cats]

        merged_keywords = ""
        try:
            merged_keywords = report_engine._resolve_category_keywords(cats, category_keywords)
        except Exception:
            merged_keywords = ""

        src_type = src.get("type", "rss")
        url_field = src.get("url", "") or ""
        direct_rss = (
            src.get("rss") or src.get("rss_url")
            or src.get("feed") or src.get("feed_url")
        )
        if not direct_rss and src_type == "rss" and str(url_field).startswith("http"):
            direct_rss = url_field

        google_rss_url = ""
        if not direct_rss:
            try:
                domain = (
                    src.get("domain") or src.get("site")
                    or report_engine._extract_news_domain(url_field) or url_field
                )
                if domain:
                    domain = str(domain).lower().replace("www.", "")
                    google_rss_url = report_engine._build_google_news_rss_for_domain(
                        domain,
                        keywords=merged_keywords,
                    )
            except Exception:
                google_rss_url = ""

        rows.append({
            "name": src.get("name", ""),
            "type": src_type,
            "category": ", ".join(cats),
            "source_url": url_field,
            "keywords_used": merged_keywords,
            "google_rss_url": google_rss_url,
            "effective_fetch_url": direct_rss or google_rss_url,
        })

    return rows


def _call_generate_report(
    start_time,
    end_time,
    selected_sources,
    selected_experts,
    profile_name,
    language,
    insights_text,
):
    fn = getattr(report_engine, "generate_report", None)
    if fn is None:
        raise RuntimeError("找不到 report_engine.generate_report()")

    sig = inspect.signature(fn)
    kwargs = {}

    candidate_values = {
        "start_time": start_time,
        "end_time": end_time,
        "selected_sources": selected_sources,
        "sources": selected_sources,
        "selected_experts": selected_experts,
        "experts": selected_experts,
        "profile_name": profile_name,
        "profile": profile_name,
        "report_template": profile_name,
        "language": language,
        "insights_text": insights_text,
        "insights": insights_text,
    }

    for name in sig.parameters:
        if name in candidate_values:
            kwargs[name] = candidate_values[name]

    result = fn(**kwargs)

    if isinstance(result, tuple):
        if len(result) >= 2:
            return result[0], result[1]
        if len(result) == 1:
            return result[0], []
    return result, []


def _fallback_save_docx(report_text: str, output_path: Path, format_config=None):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title_cfg = (format_config or {}).get("title", {})
    body_cfg = (format_config or {}).get("body", {})
    section_cfg = (format_config or {}).get("section_heading", {})

    title_size = int(title_cfg.get("font_size", 16))
    title_bold = bool(title_cfg.get("bold", True))
    title_align = title_cfg.get("align", "center")

    body_size = int(body_cfg.get("font_size", 12))
    line_spacing = float(body_cfg.get("line_spacing", 1.15))

    section_size = int(section_cfg.get("font_size", 14))
    section_bold = bool(section_cfg.get("bold", True))

    lines = (report_text or "").splitlines()

    for idx, line in enumerate(lines):
        text = line.strip()

        if idx == 0 and text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = title_bold
            run.font.size = Pt(title_size)
            if title_align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif title_align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if text.startswith("## ") or text.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(text.replace("## ", "").replace("### ", ""))
            run.bold = section_bold
            run.font.size = Pt(section_size)
            p.paragraph_format.line_spacing = line_spacing
        else:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(body_size)
            p.paragraph_format.line_spacing = line_spacing

    doc.save(str(output_path))
    return str(output_path)


def _call_save_report_docx(report_text: str, items: list, output_path: Path, format_config=None):
    fn = getattr(report_engine, "save_report_docx", None)
    if fn is None:
        return _fallback_save_docx(report_text, output_path, format_config=format_config)

    sig = inspect.signature(fn)
    params = list(sig.parameters.keys())

    try:
        if len(params) == 2:
            return fn(report_text, str(output_path))
        if len(params) == 3:
            return fn(report_text, items, str(output_path))
        if len(params) >= 4:
            return fn(report_text, items, str(output_path), format_config)
        return fn(report_text, str(output_path))
    except Exception:
        return _fallback_save_docx(report_text, output_path, format_config=format_config)


def _save_pdf_from_docx(docx_path: Path, pdf_path: Path):
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return str(pdf_path), None
    except Exception as e:
        return None, str(e)


def _try_upload_to_drive(file_path: str, folder_id: str = ""):
    if google_drive_utils is None:
        return None, "utils/google_drive.py 載入失敗"

    candidate_functions = [
        "upload_file_to_drive",
        "upload_to_drive",
        "upload_file",
    ]

    for fn_name in candidate_functions:
        fn = getattr(google_drive_utils, fn_name, None)
        if fn is None:
            continue
        try:
            sig = inspect.signature(fn)
            kwargs = {}
            if "file_path" in sig.parameters:
                kwargs["file_path"] = file_path
            elif "path" in sig.parameters:
                kwargs["path"] = file_path
            else:
                continue

            if "folder_id" in sig.parameters:
                kwargs["folder_id"] = folder_id

            result = fn(**kwargs)
            return result, None
        except Exception as e:
            return None, str(e)

    return None, "找不到可用的 Google Drive upload function"


def _format_output_targets(value):
    if isinstance(value, list):
        return value
    return [value] if value else []


def _format_output_formats(value):
    if isinstance(value, list):
        return value
    return [value] if value else []


def _parse_hhmm(text: str):
    h, m = text.split(":")
    return time(hour=int(h), minute=int(m))


def _next_daily_runs(daily_times, count=5):
    now = datetime.now()
    candidates = []
    base_date = now.date()

    for day_offset in range(0, 8):
        d = base_date + timedelta(days=day_offset)
        for t in daily_times:
            try:
                tt = _parse_hhmm(t)
                dt = datetime.combine(d, tt)
                if dt >= now:
                    candidates.append(dt)
            except Exception:
                continue

    candidates.sort()
    return candidates[:count]


def _next_interval_runs(interval_hours, window_start, window_end, count=5):
    now = datetime.now()
    interval_hours = max(1, int(interval_hours))
    candidates = []

    try:
        start_t = _parse_hhmm(window_start)
        end_t = _parse_hhmm(window_end)
    except Exception:
        start_t = time(8, 0)
        end_t = time(22, 0)

    for day_offset in range(0, 5):
        d = now.date() + timedelta(days=day_offset)
        day_start = datetime.combine(d, start_t)
        day_end = datetime.combine(d, end_t)

        current = day_start
        while current <= day_end:
            if current >= now:
                candidates.append(current)
            current += timedelta(hours=interval_hours)

    candidates.sort()
    return candidates[:count]


def _append_blank_rows(df: pd.DataFrame, blank_rows: int = 8):
    if df is None:
        df = pd.DataFrame()
    cols = list(df.columns)
    blank = pd.DataFrame([{c: "" for c in cols} for _ in range(blank_rows)])
    if "enabled" in blank.columns:
        blank["enabled"] = True
    if df.empty:
        return blank
    return pd.concat([df, blank], ignore_index=True)


def _build_source_editor_df(source_items, blank_rows=8):
    columns = ["name", "type", "url", "category", "region", "enabled", "description"]
    rows = [source_to_editor_row(x) for x in source_items]
    df = pd.DataFrame(rows, columns=columns) if rows else pd.DataFrame(columns=columns)
    df = _append_blank_rows(df, blank_rows=blank_rows)
    return df


def _build_expert_editor_df(expert_items, blank_rows=8):
    columns = ["name_zh", "name_en", "aliases", "category", "affiliation", "region", "enabled", "description"]
    rows = [expert_to_editor_row(x) for x in expert_items]
    df = pd.DataFrame(rows, columns=columns) if rows else pd.DataFrame(columns=columns)
    df = _append_blank_rows(df, blank_rows=blank_rows)
    return df


# =========================================================
# Load data
# =========================================================
all_sources = load_sources(editable_only=False)
editable_sources = [s for s in all_sources if not s.get("readonly")]
fixed_sources = [s for s in all_sources if s.get("readonly")]

experts = load_experts()
profiles = load_profiles()
formats = load_formats()
profile_map = _profiles_map(profiles)
profile_names = list(profile_map.keys()) if profile_map else ["Strategic Briefing"]
format_names = []
for f in formats:
    if isinstance(f, dict):
        n = f.get("name")
        if n and n not in format_names:
            format_names.append(n)

if "default" not in format_names:
    format_names.insert(0, "default")
saved_insights = load_insights()
auto_export_cfg = load_auto_export()

source_categories = get_source_categories(all_sources)
expert_categories = get_expert_categories(experts)

enabled_source_names = [s["name"] for s in all_sources if s.get("enabled", True)]
enabled_expert_names = [display_expert_name(e) for e in experts if e.get("enabled", True)]


# =========================================================
# Tabs
# =========================================================
tab_briefings, tab_insights, tab_sources, tab_formats, tab_automation = st.tabs(
    ["Briefings", "Insights", "Sources", "Formats", "Automation"]
)


# =========================================================
# Briefings
# =========================================================
with tab_briefings:
    st.subheader("手動生成與一鍵輸出")

    now = datetime.now()
    default_start = now - timedelta(hours=24)
    default_end = now

    c1, c2 = st.columns(2)
    with c1:
        start_date = st.date_input("開始日期", value=default_start.date(), key="brief_start_date")
        start_time_input = st.time_input(
            "開始時間",
            value=default_start.time().replace(second=0, microsecond=0),
            key="brief_start_time"
        )
    with c2:
        end_date = st.date_input("結束日期", value=default_end.date(), key="brief_end_date")
        end_time_input = st.time_input(
            "結束時間",
            value=default_end.time().replace(second=0, microsecond=0),
            key="brief_end_time"
        )

    start_dt = datetime.combine(start_date, start_time_input)
    end_dt = datetime.combine(end_date, end_time_input)

    st.markdown("---")

    c1, c2 = st.columns(2)
    with c1:
        selected_source_categories = st.multiselect(
            "來源分類",
            options=source_categories,
            default=[],
            key="brief_source_categories",
        )
        brief_source_options = [s["name"] for s in all_sources if s.get("enabled", True)]
        selected_source_names = st.multiselect(
            "來源",
            options=brief_source_options,
            default=[],
            key="brief_source_names",
        )

    with c2:
        selected_expert_categories = st.multiselect(
            "專家分類",
            options=expert_categories,
            default=[],
            key="brief_expert_categories",
        )
        selected_expert_names = st.multiselect(
            "專家",
            options=enabled_expert_names,
            default=[],
            key="brief_expert_names",
        )

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        profile_name = st.selectbox(
            "報告模板",
            options=profile_names,
            index=0,
            key="briefings_profile_selectbox",
        )
    with c2:
        selected_format_name = st.selectbox(
            "Format",
            options=format_names,
            index=format_names.index("default") if "default" in format_names else 0,
            key="briefings_format_selectbox",
        )
    with c3:
        language = st.selectbox(
            "語言",
            options=["繁體中文", "英文", "日文", "簡體中文"],
            index=0,
            key="briefings_language_selectbox",
        )
    with c4:
        output_formats = st.multiselect(
            "輸出格式",
            options=["docx", "pdf"],
            default=["docx"],
            key="briefings_output_formats",
        )

    # 讀取預設輸出設定
    default_output_targets = auto_export_cfg.get("default_output_targets", ["local"])
    default_drive_folder = auto_export_cfg.get("default_drive_folder_id", "")
    default_local_folder = auto_export_cfg.get("default_local_folder", "outputs")

    c5, c6 = st.columns(2)
    with c5:
        output_targets = st.multiselect(
            "輸出位置",
            options=["local", "google_drive"],
            default=default_output_targets,
            key="briefings_output_targets",
        )
    with c6:

        local_folder = st.text_input(
            "Local 輸出資料夾",
            value=default_local_folder,
            key="briefings_local_folder",
        )

        google_drive_folder_id = st.text_input(
            "Google Drive Folder ID",
            value=default_drive_folder,
            key="briefings_drive_folder_id",
        )

    extra_insights = st.text_area(
        "本次補充 insights",
        value="",
        height=180,
        key="briefings_extra_insights",
    )

    selected_sources = _filter_sources_by_selection(
        all_sources,
        selected_source_categories,
        selected_source_names,
    )
    selected_experts = _filter_experts_by_selection(
        experts,
        selected_expert_categories,
        selected_expert_names,
    )

    s1, s2, s3 = st.columns(3)
    s1.metric("已選來源數", len(selected_sources))
    s2.metric("已選專家數", len(selected_experts))
    s3.metric("時間範圍（小時）", round(max((end_dt - start_dt).total_seconds() / 3600, 0), 2))

    if st.button("一鍵生成並輸出", type="primary", use_container_width=True, key="run_briefings"):
        if start_dt >= end_dt:
            st.error("開始時間必須早於結束時間。")
        elif not output_formats:
            st.error("請至少選一種輸出格式。")
        elif not output_targets:
            st.error("請至少選一種輸出位置。")
        else:
            status = st.empty()
            detail = st.empty()
            progress = st.progress(0)
            result_box = st.container()

            try:
                status.info("開始生成")
                detail.caption("正在準備任務參數。")
                progress.progress(10)

                insight_lines = []

                if isinstance(saved_insights, list):
                    for ins in saved_insights:
                        if not isinstance(ins, dict):
                            continue

                        title = str(ins.get("title", "")).strip()
                        content = str(ins.get("content", "")).strip()

                        if title and content:
                            insight_lines.append(f"{title}: {content}")
                        elif content:
                            insight_lines.append(content)

                elif isinstance(saved_insights, str) and saved_insights.strip():
                    insight_lines.append(saved_insights.strip())

                combined_insights = "\n".join(insight_lines)
                selected_format_config = next(
                    (f for f in formats if f.get("name") == selected_format_name),
                    next((f for f in formats if f.get("name") == "default"), None)
                )
                if extra_insights.strip():
                    combined_insights = f"{combined_insights}\n\n{extra_insights.strip()}".strip()

                status.info("抓取資料並生成 AI 簡報")
                detail.markdown(
                    f"""
**本次時間範圍：** {start_dt.strftime('%Y-%m-%d %H:%M')} ～ {end_dt.strftime('%Y-%m-%d %H:%M')}

**已選來源數：** {len(selected_sources)}

**已選專家數：** {len(selected_experts)}
"""
                )
                progress.progress(45)

                report_text, filtered_items = _call_generate_report(
                    start_time=start_dt,
                    end_time=end_dt,
                    selected_sources=selected_sources,
                    selected_experts=selected_experts,
                    profile_name=profile_name,
                    language=language,
                    insights_text=combined_insights,
                )

                status.info("輸出檔案")
                progress.progress(75)

                ts = _now_str()
                base_name = f"briefings_{ts}"
                output_dir = Path(local_folder)
                output_dir.mkdir(parents=True, exist_ok=True)

                docx_path = output_dir / f"{base_name}.docx"
                pdf_path = output_dir / f"{base_name}.pdf"

                output_logs = []
                generated_files = []

                if "docx" in output_formats or "pdf" in output_formats:
                    _call_save_report_docx(report_text, filtered_items, docx_path, selected_format_config)
                    generated_files.append(docx_path)

                if "pdf" in output_formats:
                    pdf_result, pdf_error = _save_pdf_from_docx(docx_path, pdf_path)
                    if pdf_result:
                        generated_files.append(pdf_path)
                    else:
                        output_logs.append(f"PDF 轉換失敗：{pdf_error}")

                if "local" in output_targets:
                    for f in generated_files:
                        output_logs.append(f"已輸出到本機：{f}")

                if "google_drive" in output_targets:
                    for f in generated_files:
                        uploaded, err = _try_upload_to_drive(str(f), google_drive_folder_id)
                        if err:
                            output_logs.append(f"Google Drive 上傳失敗（{f.name}）：{err}")
                        else:
                            output_logs.append(f"Google Drive 已上傳：{f.name}")

                progress.progress(100)
                status.success("完成")
                # 儲存預設輸出設定
                auto_export_cfg["default_output_targets"] = output_targets
                auto_export_cfg["default_drive_folder_id"] = google_drive_folder_id
                auto_export_cfg["default_local_folder"] = local_folder

                save_auto_export(auto_export_cfg)

                with result_box:
                    st.markdown("### 執行結果")
                    st.write(f"**本次時間範圍：** {start_dt} ～ {end_dt}")
                    st.write(f"**已選來源數：** {len(selected_sources)}")
                    st.write(f"**已選專家數：** {len(selected_experts)}")
                    st.write(f"**已取得文章數：** {len(filtered_items) if filtered_items is not None else 0}")

                    if output_logs:
                        st.markdown("**輸出結果：**")
                        for line in output_logs:
                            st.write(f"- {line}")

                    st.markdown("### 報告預覽")
                    st.text_area("report", value=report_text or "", height=400, label_visibility="collapsed", key="briefings_report_preview")

            except Exception as e:
                status.error(f"生成失敗：{e}")


# =========================================================
# Insights
# =========================================================
with tab_insights:

    st.subheader("Insights Editor")
    st.caption("Strategic guidance used by AI when generating reports. These will NOT be cited.")

    insights_data = load_insights()

    # 相容舊版：如果還是 str，就先轉成空列表
    if isinstance(insights_data, str):
        insights_data = []
    elif not isinstance(insights_data, list):
        insights_data = []

    table_rows = []
    for item in insights_data:
        if not isinstance(item, dict):
            continue

        tags_value = item.get("tags", [])
        if isinstance(tags_value, list):
            tags_str = ", ".join([str(x).strip() for x in tags_value if str(x).strip()])
        elif isinstance(tags_value, str):
            tags_str = tags_value
        else:
            tags_str = ""

        table_rows.append({
            "title": item.get("title", ""),
            "content": item.get("content", ""),
            "tags": tags_str,
        })

    df = pd.DataFrame(table_rows, columns=["title", "content", "tags"])

    # 新增空白列讓 editor 可新增
    blank_rows = pd.DataFrame([
        {"title": "", "content": "", "tags": ""}
        for _ in range(10)
    ])

    df = pd.concat([df, blank_rows], ignore_index=True)

    edited_df = st.data_editor(
        df,
        num_rows="dynamic",
        use_container_width=True,
        height=420,
        key="insights_editor",
        column_config={
            "title": st.column_config.TextColumn(
                "Title",
                width="medium",
                help="Short headline for this insight"
            ),
            "content": st.column_config.TextColumn(
                "Content",
                width="large",
                help="Strategic guidance for the AI"
            ),
            "tags": st.column_config.TextColumn(
            "Tags",
            width="medium",
            help="Comma separated topics (optional)"
        ),
    },
)

    if st.button("Save Insights", key="save_insights_btn", use_container_width=True):
        cleaned = []

        for _, row in edited_df.iterrows():
            title = str(row.get("title", "")).strip()
            content = str(row.get("content", "")).strip()
            tags_raw = str(row.get("tags", "")).strip()

            if not title and not content and not tags_raw:
                continue

            tag_list = [t.strip() for t in tags_raw.split(",") if t.strip()]

            cleaned.append({
                "title": title,
                "content": content,
                "tags": tag_list,
            })

        save_insights(cleaned)
        st.success("Insights saved.")
        st.rerun()


# =========================================================
# Sources
# =========================================================
with tab_sources:
    st.subheader("Sources 管理")

    with st.expander("單筆新增來源", expanded=False):
        c1, c2 = st.columns(2)
        with c1:
            src_name = st.text_input("name", key="single_src_name")
            src_type = st.selectbox("type", options=["rss", "domain"], key="single_src_type")
            src_url = st.text_input("url", key="single_src_url")
            src_category = st.text_input("category（可輸入多個，以逗號分隔）", key="single_src_category")
        with c2:
            src_region = st.text_input("region", key="single_src_region")
            src_enabled = st.checkbox("enabled", value=True, key="single_src_enabled")
            src_description = st.text_area("description", key="single_src_description", height=120)

        if st.button("新增來源", key="add_single_source"):
            new_item = editor_row_to_source(
                {
                    "name": src_name,
                    "type": src_type,
                    "url": src_url,
                    "category": src_category,
                    "region": src_region,
                    "enabled": src_enabled,
                    "description": src_description,
                }
            )
            current = load_sources(editable_only=True)
            if not new_item["name"]:
                st.error("來源名稱不可空白。")
            else:
                current.append(new_item)
                save_sources(current)
                st.success("已新增來源。")
                st.rerun()

    st.markdown("### 表格式批次貼上新增來源")
    st.caption("可直接從外部複製多列資料貼到下表，再按「批次加入來源」。")

    source_batch_columns = ["name", "type", "url", "category", "region", "enabled", "description"]
    source_batch_default = pd.DataFrame([{c: "" for c in source_batch_columns} for _ in range(8)])
    source_batch_default["enabled"] = True

    source_batch_df = st.data_editor(
        source_batch_default,
        num_rows="dynamic",
        use_container_width=True,
        height=280,
        key="source_batch_editor",
        column_config={
            "name": st.column_config.TextColumn("name"),
            "type": st.column_config.SelectboxColumn("type", options=["rss", "domain"]),
            "url": st.column_config.TextColumn("url"),
            "category": st.column_config.TextColumn("category"),
            "region": st.column_config.TextColumn("region"),
            "enabled": st.column_config.CheckboxColumn("enabled", default=True),
            "description": st.column_config.TextColumn("description"),
        },
    )

    if st.button("批次加入來源", key="batch_add_sources"):
        rows = _clean_batch_df(source_batch_df)
        if not rows:
            st.warning("沒有可加入的來源資料。")
        else:
            current = load_sources(editable_only=True)
            name_set = {x.get("name", "").strip() for x in current}
            added = 0
            for row in rows:
                item = editor_row_to_source(row)
                if not item["name"]:
                    continue
                if item["name"] in name_set:
                    current = [x for x in current if x.get("name") != item["name"]]
                current.append(item)
                name_set.add(item["name"])
                added += 1
            save_sources(current)
            st.success(f"已批次加入 / 更新 {added} 筆來源。")
            st.rerun()

    st.markdown("### 既有可編輯來源清單")
    st.caption("這裡也可以直接貼上多列資料、修改既有資料、增加新列，再按儲存。")

    editable_source_df = _build_source_editor_df(editable_sources, blank_rows=10)

    edited_sources_df = st.data_editor(
        editable_source_df,
        num_rows="dynamic",
        use_container_width=True,
        height=420,
        key="editable_sources_editor",
        column_config={
            "name": st.column_config.TextColumn("name"),
            "type": st.column_config.SelectboxColumn("type", options=["rss", "domain"]),
            "url": st.column_config.TextColumn("url"),
            "category": st.column_config.TextColumn("category"),
            "region": st.column_config.TextColumn("region"),
            "enabled": st.column_config.CheckboxColumn("enabled", default=True),
            "description": st.column_config.TextColumn("description"),
        },
    )

    c1, c2 = st.columns(2)
    with c1:
        if st.button("儲存來源清單編輯", key="save_sources_table", use_container_width=True):
            rows = _clean_batch_df(edited_sources_df)
            cleaned = []
            for row in rows:
                item = editor_row_to_source(row)
                if item["name"]:
                    cleaned.append(item)
            save_sources(cleaned)
            st.success("來源清單已儲存。")
            st.rerun()

    with c2:
        source_delete_options = [s["name"] for s in editable_sources]
        delete_source_names = st.multiselect("刪除來源", options=source_delete_options, key="delete_source_names")
        if st.button("刪除選取來源", key="delete_sources_btn", use_container_width=True):
            current = load_sources(editable_only=True)
            current = [x for x in current if x.get("name") not in delete_source_names]
            save_sources(current)
            st.success(f"已刪除 {len(delete_source_names)} 筆來源。")
            st.rerun()

    st.markdown("### 固定中共官媒來源（唯讀）")
    fixed_rows = [source_to_editor_row(x) for x in fixed_sources]
    fixed_df = pd.DataFrame(fixed_rows) if fixed_rows else pd.DataFrame(columns=source_batch_columns)
    st.dataframe(fixed_df, use_container_width=True, hide_index=True)

    st.markdown("### 實際抓取 URL 預覽")
    st.caption("domain 類型會顯示實際生成的 Google News RSS URL，並納入此來源所有 category 對應的關鍵字。")
    preview_rows = _build_source_fetch_preview_rows(load_sources())
    preview_df = pd.DataFrame(preview_rows)
    if not preview_df.empty:
        st.dataframe(preview_df, use_container_width=True, hide_index=True)

    st.markdown("---")
    st.subheader("Experts 管理（仍在 Sources 頁下方）")

    with st.expander("單筆新增專家", expanded=False):
        c1, c2 = st.columns(2)
        with c1:
            exp_name_zh = st.text_input("中文名 name_zh", key="single_exp_name_zh")
            exp_name_en = st.text_input("英文名 name_en", key="single_exp_name_en")
            exp_aliases = st.text_input("aliases（逗號分隔）", key="single_exp_aliases")
            exp_category = st.text_input("category（逗號分隔）", key="single_exp_category")
        with c2:
            exp_affiliation = st.text_input("affiliation", key="single_exp_affiliation")
            exp_region = st.text_input("region", key="single_exp_region")
            exp_enabled = st.checkbox("enabled", value=True, key="single_exp_enabled")
            exp_description = st.text_area("description", key="single_exp_description", height=120)

        if st.button("新增專家", key="add_single_expert"):
            new_item = editor_row_to_expert(
                {
                    "name_zh": exp_name_zh,
                    "name_en": exp_name_en,
                    "aliases": exp_aliases,
                    "category": exp_category,
                    "affiliation": exp_affiliation,
                    "region": exp_region,
                    "enabled": exp_enabled,
                    "description": exp_description,
                }
            )
            current = load_experts()
            display_name = display_expert_name(new_item)
            if not display_name or display_name == "Unnamed Expert":
                st.error("至少要填中文名或英文名。")
            else:
                current.append(new_item)
                save_experts(current)
                st.success("已新增專家。")
                st.rerun()

    st.markdown("### 表格式批次貼上新增專家")
    st.caption("可直接從外部複製多列資料貼到下表，再按「批次加入專家」。")

    expert_batch_columns = ["name_zh", "name_en", "aliases", "category", "affiliation", "region", "enabled", "description"]
    expert_batch_default = pd.DataFrame([{c: "" for c in expert_batch_columns} for _ in range(8)])
    expert_batch_default["enabled"] = True

    expert_batch_df = st.data_editor(
        expert_batch_default,
        num_rows="dynamic",
        use_container_width=True,
        height=280,
        key="expert_batch_editor",
        column_config={
            "name_zh": st.column_config.TextColumn("name_zh"),
            "name_en": st.column_config.TextColumn("name_en"),
            "aliases": st.column_config.TextColumn("aliases"),
            "category": st.column_config.TextColumn("category"),
            "affiliation": st.column_config.TextColumn("affiliation"),
            "region": st.column_config.TextColumn("region"),
            "enabled": st.column_config.CheckboxColumn("enabled", default=True),
            "description": st.column_config.TextColumn("description"),
        },
    )

    if st.button("批次加入專家", key="batch_add_experts"):
        rows = _clean_batch_df(expert_batch_df)
        if not rows:
            st.warning("沒有可加入的專家資料。")
        else:
            current = load_experts()
            name_set = {display_expert_name(x) for x in current}
            added = 0
            for row in rows:
                item = editor_row_to_expert(row)
                disp = display_expert_name(item)
                if not disp or disp == "Unnamed Expert":
                    continue
                if disp in name_set:
                    current = [x for x in current if display_expert_name(x) != disp]
                current.append(item)
                name_set.add(disp)
                added += 1
            save_experts(current)
            st.success(f"已批次加入 / 更新 {added} 筆專家。")
            st.rerun()

    st.markdown("### 既有專家清單")
    st.caption("這裡也可以直接貼上多列資料、修改既有資料、增加新列，再按儲存。")

    experts_df = _build_expert_editor_df(experts, blank_rows=10)

    edited_experts_df = st.data_editor(
        experts_df,
        num_rows="dynamic",
        use_container_width=True,
        height=420,
        key="editable_experts_editor",
        column_config={
            "name_zh": st.column_config.TextColumn("name_zh"),
            "name_en": st.column_config.TextColumn("name_en"),
            "aliases": st.column_config.TextColumn("aliases"),
            "category": st.column_config.TextColumn("category"),
            "affiliation": st.column_config.TextColumn("affiliation"),
            "region": st.column_config.TextColumn("region"),
            "enabled": st.column_config.CheckboxColumn("enabled", default=True),
            "description": st.column_config.TextColumn("description"),
        },
    )

    c1, c2 = st.columns(2)
    with c1:
        if st.button("儲存專家清單編輯", key="save_experts_table", use_container_width=True):
            rows = _clean_batch_df(edited_experts_df)
            cleaned = []
            for row in rows:
                item = editor_row_to_expert(row)
                disp = display_expert_name(item)
                if disp and disp != "Unnamed Expert":
                    cleaned.append(item)
            save_experts(cleaned)
            st.success("專家清單已儲存。")
            st.rerun()

    with c2:
        expert_delete_options = [display_expert_name(x) for x in experts]
        delete_expert_names = st.multiselect("刪除專家", options=expert_delete_options, key="delete_expert_names")
        if st.button("刪除選取專家", key="delete_experts_btn", use_container_width=True):
            current = load_experts()
            current = [x for x in current if display_expert_name(x) not in delete_expert_names]
            save_experts(current)
            st.success(f"已刪除 {len(delete_expert_names)} 筆專家。")
            st.rerun()

    # ── 各類別專家 Google News RSS 關鍵字設定 ──────────────────────────────────
    st.markdown("### 專家 Google News RSS 關鍵字篩選")
    st.caption("各類別專家在搜尋時，會將人名與下方關鍵字組合（人名 AND 關鍵字），讓結果更聚焦於相關主題。")

    _exp_cat_kw = load_category_keywords()

    exp_kw_tabs = st.tabs(["🇹🇼 台灣專家", "🌐 國際專家", "🇨🇳 中國專家"])

    with exp_kw_tabs[0]:
        with st.expander("🔍 Google News RSS 關鍵字篩選", expanded=False):
            _kw_tw_exp = st.text_area(
                "台灣專家 關鍵字",
                value=_exp_cat_kw.get("台灣專家", DEFAULT_CATEGORY_KEYWORDS.get("台灣專家", "")),
                height=100,
                key="kw_tw_expert",
            )
            if st.button("儲存台灣專家關鍵字", key="save_kw_tw_expert"):
                _exp_cat_kw["台灣專家"] = _kw_tw_exp.strip()
                save_category_keywords(_exp_cat_kw)
                st.success("已儲存台灣專家關鍵字。")

    with exp_kw_tabs[1]:
        with st.expander("🔍 Google News RSS 關鍵字篩選", expanded=False):
            _kw_intl_exp = st.text_area(
                "國際專家 關鍵字",
                value=_exp_cat_kw.get("國際專家", DEFAULT_CATEGORY_KEYWORDS.get("國際專家", "")),
                height=100,
                key="kw_intl_expert",
            )
            if st.button("儲存國際專家關鍵字", key="save_kw_intl_expert"):
                _exp_cat_kw["國際專家"] = _kw_intl_exp.strip()
                save_category_keywords(_exp_cat_kw)
                st.success("已儲存國際專家關鍵字。")

    with exp_kw_tabs[2]:
        with st.expander("🔍 Google News RSS 關鍵字篩選", expanded=False):
            _kw_cn_exp = st.text_area(
                "中國專家 關鍵字",
                value=_exp_cat_kw.get("中國專家", DEFAULT_CATEGORY_KEYWORDS.get("中國專家", "")),
                height=100,
                key="kw_cn_expert",
            )
            if st.button("儲存中國專家關鍵字", key="save_kw_cn_expert"):
                _exp_cat_kw["中國專家"] = _kw_cn_exp.strip()
                save_category_keywords(_exp_cat_kw)
                st.success("已儲存中國專家關鍵字。")

    st.markdown("### 專家搜尋名稱預覽")
    preview_rows = []
    for e in load_experts():
        preview_rows.append(
            {
                "display_name": display_expert_name(e),
                "search_names": ", ".join(e.get("search_names", [])),
                "aliases": ", ".join(e.get("aliases", [])),
                "category": ", ".join(e.get("category", [])),
                "affiliation": e.get("affiliation", ""),
                "region": e.get("region", ""),
                "enabled": e.get("enabled", True),
            }
        )
    if preview_rows:
        st.dataframe(pd.DataFrame(preview_rows), use_container_width=True, hide_index=True)
    else:
        st.info("目前尚無專家資料。")


# =========================================================
# Formats
# =========================================================
with tab_formats:
    st.subheader("Report Formats")

    formats = load_formats()
    if not formats:
        st.info("目前沒有 formats。請先建立 config/formats.json")
    else:
        selected_format = st.selectbox(
            "選擇 Format",
            options=[f.get("name", "default") for f in formats],
            index=0,
            key="formats_selected_name",
        )

        current = next(
            (f for f in formats if f.get("name", "default") == selected_format),
            None
        )

        if current:
            c1, c2 = st.columns(2)

            with c1:
                title_font_size = st.number_input(
                    "Title Font Size",
                    min_value=8,
                    max_value=48,
                    value=int(current.get("title", {}).get("font_size", 16)),
                    step=1,
                )
                title_bold = st.checkbox(
                    "Title Bold",
                    value=bool(current.get("title", {}).get("bold", True)),
                )

                section_font_size = st.number_input(
                    "Section Heading Font Size",
                    min_value=8,
                    max_value=36,
                    value=int(current.get("section_heading", {}).get("font_size", 14)),
                    step=1,
                )
                section_bold = st.checkbox(
                    "Section Heading Bold",
                    value=bool(current.get("section_heading", {}).get("bold", True)),
                )

            with c2:
                body_font_size = st.number_input(
                    "Body Font Size",
                    min_value=8,
                    max_value=24,
                    value=int(current.get("body", {}).get("font_size", 12)),
                    step=1,
                )
                line_spacing = st.number_input(
                    "Line Spacing",
                    min_value=1.0,
                    max_value=3.0,
                    value=float(current.get("body", {}).get("line_spacing", 1.15)),
                    step=0.05,
                )

                notes_style = st.selectbox(
                    "Notes Style",
                    options=["footnote", "endnote", "none"],
                    index=["footnote", "endnote", "none"].index(
                        current.get("notes", {}).get("style", "footnote")
                    ) if current.get("notes", {}).get("style", "footnote") in ["footnote", "endnote", "none"] else 0,
                )

                link_placement = st.selectbox(
                    "Link Placement",
                    options=["inline", "footnote", "none"],
                    index=["inline", "footnote", "none"].index(
                        current.get("links", {}).get("placement", "inline")
                    ) if current.get("links", {}).get("placement", "inline") in ["inline", "footnote", "none"] else 0,
                )

            if st.button("Save Format", key="save_format_btn", use_container_width=True):
                current.setdefault("title", {})
                current.setdefault("section_heading", {})
                current.setdefault("body", {})
                current.setdefault("notes", {})
                current.setdefault("links", {})

                current["title"]["font_size"] = int(title_font_size)
                current["title"]["bold"] = bool(title_bold)

                current["section_heading"]["font_size"] = int(section_font_size)
                current["section_heading"]["bold"] = bool(section_bold)

                current["body"]["font_size"] = int(body_font_size)
                current["body"]["line_spacing"] = float(line_spacing)

                current["notes"]["style"] = notes_style
                current["links"]["placement"] = link_placement

                save_formats(formats)
                st.success("Format 已儲存。")

# =========================================================
# Automation
# =========================================================
# =========================================================
# Automation
# =========================================================
with tab_automation:
    import pandas as pd
    from datetime import datetime
    from utils.auto_export import compute_next_runs

    st.subheader("AI Briefings Scheduler")

    config = load_auto_export()
    config.setdefault("enabled", True)
    config.setdefault("schedules", [])

    schedules = config["schedules"]

    # -------------------------
    # 載入來源 / 專家 / 模板
    # -------------------------
    all_sources = load_sources()
    all_experts = load_experts()
    all_profiles = load_profiles()

    source_categories = get_source_categories(all_sources)
    expert_categories = get_expert_categories(all_experts)

    profile_names = []
    for p in all_profiles:
        if isinstance(p, dict):
            name = (p.get("name") or "").strip()
            if name:
                profile_names.append(name)
    if "default" not in profile_names:
        profile_names = ["default"] + profile_names

    source_name_options = []
    for s in all_sources:
        n = (s.get("name") or "").strip()
        if n and n not in source_name_options:
            source_name_options.append(n)

    expert_name_options = []
    for e in all_experts:
        n = display_expert_name(e)
        if n and n not in expert_name_options:
            expert_name_options.append(n)

    # -------------------------
    # helpers
    # -------------------------
    def _csv_to_list(text):
        if text is None:
            return []
        text = str(text).strip()
        if not text:
            return []
        for sep in ["\n", "；", ";", "、"]:
            text = text.replace(sep, ",")
        return [x.strip() for x in text.split(",") if x.strip()]

    def _safe_int(v, default):
        try:
            return int(v)
        except Exception:
            return default

    def _schedule_to_table_row(s):
        mode = s.get("schedule_mode", "daily")

        if mode == "once":
            time_or_interval = s.get("once_datetime", "")
        elif mode == "hourly":
            time_or_interval = f'{s.get("hourly_interval_hours", 1)}h'
        elif mode == "daily":
            time_or_interval = ", ".join(s.get("daily_times", []))
        elif mode == "weekly":
            days = ",".join([str(x) for x in s.get("weekly_days", [])])
            times = ",".join(s.get("weekly_times", []))
            time_or_interval = f"{days} @ {times}"
        elif mode == "monthly":
            days = ",".join([str(x) for x in s.get("monthly_days", [])])
            times = ",".join(s.get("monthly_times", []))
            time_or_interval = f"{days} @ {times}"
        else:
            time_or_interval = ""

        return {
            "name": s.get("name", ""),
            "schedule_mode": mode,
            "time_or_interval": time_or_interval,
            "language": s.get("language", "繁體中文"),
            "profile": s.get("profile", "default"),
            "format_name": s.get("format_name", "default"),
            "output_formats": ", ".join(s.get("output_formats", ["docx"])),
            "output_targets": ", ".join(s.get("output_targets", ["local"])),
            "coverage_hours": s.get("coverage_hours", 24),
            "delete": False,
        }

    def _table_row_to_schedule(row, base_schedule):
        s = dict(base_schedule)

        s["name"] = str(row.get("name", "")).strip() or "Untitled Schedule"
        s["schedule_mode"] = str(row.get("schedule_mode", "daily")).strip() or "daily"
        s["language"] = str(row.get("language", "繁體中文")).strip() or "繁體中文"
        s["profile"] = str(row.get("profile", "default")).strip() or "default"
        s["format_name"] = str(row.get("format_name", "default")).strip() if "format_name" in row else "default"
        s["output_formats"] = _csv_to_list(row.get("output_formats", "docx")) or ["docx"]
        s["output_targets"] = _csv_to_list(row.get("output_targets", "local")) or ["local"]
        s["coverage_hours"] = max(1, _safe_int(row.get("coverage_hours", 24), 24))

        mode = s["schedule_mode"]
        raw = str(row.get("time_or_interval", "")).strip()

        # 先保留舊值，避免編輯表格時誤清空其他模式欄位
        s.setdefault("once_datetime", "")
        s.setdefault("hourly_interval_hours", 1)
        s.setdefault("daily_times", ["09:00"])
        s.setdefault("weekly_days", [0])
        s.setdefault("weekly_times", ["09:00"])
        s.setdefault("monthly_days", [1])
        s.setdefault("monthly_times", ["09:00"])

        if mode == "once":
            s["once_datetime"] = raw

        elif mode == "hourly":
            txt = raw.lower().replace("hours", "").replace("hour", "").replace("h", "").strip()
            s["hourly_interval_hours"] = max(1, _safe_int(txt, s.get("hourly_interval_hours", 1)))

        elif mode == "daily":
            s["daily_times"] = _csv_to_list(raw) or s.get("daily_times", ["09:00"])

        elif mode == "weekly":
            # 格式: 1,3,5 @ 09:00,18:00
            if "@" in raw:
                left, right = raw.split("@", 1)
                weekly_days = []
                for x in _csv_to_list(left):
                    try:
                        d = int(x)
                        if 0 <= d <= 6:
                            weekly_days.append(d)
                    except Exception:
                        pass
                weekly_times = _csv_to_list(right)
                s["weekly_days"] = weekly_days or s.get("weekly_days", [0])
                s["weekly_times"] = weekly_times or s.get("weekly_times", ["09:00"])

        elif mode == "monthly":
            # 格式: 1,15,28 @ 09:00,18:00
            if "@" in raw:
                left, right = raw.split("@", 1)
                monthly_days = []
                for x in _csv_to_list(left):
                    try:
                        d = int(x)
                        if 1 <= d <= 31:
                            monthly_days.append(d)
                    except Exception:
                        pass
                monthly_times = _csv_to_list(right)
                s["monthly_days"] = monthly_days or s.get("monthly_days", [1])
                s["monthly_times"] = monthly_times or s.get("monthly_times", ["09:00"])

        return s

    def _new_schedule():
        return {
            "name": f"Briefing {len(config['schedules']) + 1}",
            "schedule_mode": "daily",
            "once_datetime": "",
            "hourly_interval_hours": 4,
            "daily_times": ["09:00"],
            "weekly_days": [0],
            "weekly_times": ["09:00"],
            "monthly_days": [1],
            "monthly_times": ["09:00"],
            "coverage_hours": 24,
            "selected_source_categories": [],
            "selected_source_names": [],
            "selected_expert_categories": [],
            "selected_expert_names": [],
            "language": "繁體中文",
            "profile": "default",
            "format_name": "default",
            "output_formats": ["docx"],
            "output_targets": ["local"],
            "google_drive_folder_id": "",
        }

    if "automation_selected_index" not in st.session_state:
        st.session_state.automation_selected_index = 0

    # -------------------------
    # 啟用開關
    # -------------------------
    top_c1, top_c2, top_c3 = st.columns([1, 1, 1])

    with top_c1:
        config["enabled"] = st.checkbox("啟用自動排程", value=config.get("enabled", True))

    with top_c2:
        if st.button("➕ 新增排程", use_container_width=True):
            config["schedules"].append(_new_schedule())
            save_auto_export(config)
            st.session_state.automation_selected_index = max(0, len(config["schedules"]) - 1)
            st.rerun()

    with top_c3:
        if st.button("💾 儲存全部排程", use_container_width=True):
            save_auto_export(config)
            st.success("已儲存排程設定")

    # -------------------------
    # 排程表
    # -------------------------
    st.markdown("### 排程列表")

    table_rows = [_schedule_to_table_row(s) for s in schedules]
    if not table_rows:
        table_rows = [{
            "name": "",
            "schedule_mode": "daily",
            "time_or_interval": "",
            "language": "繁體中文",
            "profile": "default",
            "format_name": "default",
            "output_formats": "docx",
            "output_targets": "local",
            "coverage_hours": 24,
            "delete": False,
        }]

    table_df = pd.DataFrame(table_rows)

    edited_df = st.data_editor(
        table_df,
        use_container_width=True,
        num_rows="fixed",
        column_config={
            "name": st.column_config.TextColumn("排程名稱"),
            "schedule_mode": st.column_config.SelectboxColumn(
                "模式",
                options=["once", "hourly", "daily", "weekly", "monthly"],
            ),
            "time_or_interval": st.column_config.TextColumn("時間 / 間隔"),
            "language": st.column_config.SelectboxColumn(
                "語言",
                options=["繁體中文", "英文", "日文", "簡體中文"],
            ),
            "profile": st.column_config.SelectboxColumn(
                "Profile",
                options=profile_names,
            ),
            "format_name": st.column_config.SelectboxColumn(
                "Format",
                options=format_names,
            ),
            "output_formats": st.column_config.TextColumn("格式"),
            "output_targets": st.column_config.TextColumn("輸出位置"),
            "coverage_hours": st.column_config.NumberColumn("涵蓋小時", min_value=1, step=1),
            "delete": st.column_config.CheckboxColumn("刪除"),
        },
        key="automation_schedule_editor",
    )

    if len(schedules) > 0:
        rebuilt = []
        for i, row in edited_df.iterrows():
            if bool(row.get("delete", False)):
                continue
            base_schedule = schedules[i] if i < len(schedules) else _new_schedule()
            rebuilt.append(_table_row_to_schedule(dict(row), base_schedule))
        config["schedules"] = rebuilt
        schedules = config["schedules"]

        if st.session_state.automation_selected_index >= len(schedules):
            st.session_state.automation_selected_index = max(0, len(schedules) - 1)

    # -------------------------
    # 細部設定
    # -------------------------
    st.markdown("### 排程細部設定")

    if not schedules:
        st.info("目前沒有排程。請先按「新增排程」。")
    else:
        selected_options = [
            f"{idx + 1}. {s.get('name', 'Untitled Schedule')}"
            for idx, s in enumerate(schedules)
        ]

        selected_label = st.selectbox(
            "選擇排程",
            options=selected_options,
            index=min(st.session_state.automation_selected_index, len(selected_options) - 1),
        )
        selected_idx = selected_options.index(selected_label)
        st.session_state.automation_selected_index = selected_idx

        s = schedules[selected_idx]

        c1, c2 = st.columns(2)

        with c1:
            s["name"] = st.text_input("排程名稱", value=s.get("name", ""), key=f"name_{selected_idx}")
            s["coverage_hours"] = int(
                st.number_input(
                    "涵蓋過去幾小時",
                    min_value=1,
                    max_value=720,
                    value=int(s.get("coverage_hours", 24)),
                    step=1,
                    key=f"coverage_{selected_idx}",
                )
            )
            s["language"] = st.selectbox(
                "語言",
                options=["繁體中文", "英文", "日文", "簡體中文"],
                index=["繁體中文", "英文", "日文", "簡體中文"].index(
                    s.get("language", "繁體中文")
                ) if s.get("language", "繁體中文") in ["繁體中文", "英文", "日文", "簡體中文"] else 0,
                key=f"language_{selected_idx}",
            )
            s["profile"] = st.selectbox(
                "模板",
                options=profile_names,
                index=profile_names.index(s.get("profile", "default"))
                if s.get("profile", "default") in profile_names else 0,
                key=f"profile_{selected_idx}",
            )

            s["format_name"] = st.selectbox(
                "Format",
                options=format_names,
                index=format_names.index(s.get("format_name", "default"))
                if s.get("format_name", "default") in format_names else 0,
                key=f"format_{selected_idx}",
            )
            s["output_formats"] = st.multiselect(
                "輸出格式",
                options=["docx", "pdf"],
                default=s.get("output_formats", ["docx"]),
                key=f"formats_{selected_idx}",
            )
            s["output_targets"] = st.multiselect(
                "輸出位置",
                options=["local", "google_drive"],
                default=s.get("output_targets", ["local"]),
                key=f"targets_{selected_idx}",
            )
            s["google_drive_folder_id"] = st.text_input(
                "Google Drive Folder ID",
                value=s.get("google_drive_folder_id", ""),
                key=f"gdrive_{selected_idx}",
            )

        with c2:
            s["schedule_mode"] = st.selectbox(
                "排程模式",
                options=["once", "hourly", "daily", "weekly", "monthly"],
                index=["once", "hourly", "daily", "weekly", "monthly"].index(
                    s.get("schedule_mode", "daily")
                ),
                key=f"mode_{selected_idx}",
            )

            mode = s["schedule_mode"]

            if mode == "once":
                s["once_datetime"] = st.text_input(
                    "指定時間（YYYY-MM-DD HH:MM）",
                    value=s.get("once_datetime", datetime.now().strftime("%Y-%m-%d %H:%M")),
                    key=f"once_{selected_idx}",
                )

            elif mode == "hourly":
                s["hourly_interval_hours"] = int(
                    st.number_input(
                        "每幾小時執行一次",
                        min_value=1,
                        max_value=168,
                        value=int(s.get("hourly_interval_hours", 4)),
                        step=1,
                        key=f"hourly_{selected_idx}",
                    )
                )

            elif mode == "daily":
                daily_text = st.text_input(
                    "每日時間（可多個，以逗號分隔）",
                    value=", ".join(s.get("daily_times", ["09:00"])),
                    key=f"daily_{selected_idx}",
                )
                s["daily_times"] = _csv_to_list(daily_text) or ["09:00"]

            elif mode == "weekly":
                weekly_days = st.multiselect(
                    "每週星期幾（0=Mon, 6=Sun）",
                    options=[0, 1, 2, 3, 4, 5, 6],
                    default=s.get("weekly_days", [0]),
                    key=f"weekly_days_{selected_idx}",
                )
                weekly_times_text = st.text_input(
                    "每週時間（可多個，以逗號分隔）",
                    value=", ".join(s.get("weekly_times", ["09:00"])),
                    key=f"weekly_times_{selected_idx}",
                )
                s["weekly_days"] = weekly_days or [0]
                s["weekly_times"] = _csv_to_list(weekly_times_text) or ["09:00"]

            elif mode == "monthly":
                monthly_days_text = st.text_input(
                    "每月幾號（可多個，以逗號分隔）",
                    value=", ".join([str(x) for x in s.get("monthly_days", [1])]),
                    key=f"monthly_days_{selected_idx}",
                )
                monthly_times_text = st.text_input(
                    "每月時間（可多個，以逗號分隔）",
                    value=", ".join(s.get("monthly_times", ["09:00"])),
                    key=f"monthly_times_{selected_idx}",
                )
                month_days = []
                for x in _csv_to_list(monthly_days_text):
                    try:
                        d = int(x)
                        if 1 <= d <= 31:
                            month_days.append(d)
                    except Exception:
                        pass
                s["monthly_days"] = month_days or [1]
                s["monthly_times"] = _csv_to_list(monthly_times_text) or ["09:00"]

        st.markdown("#### 來源與專家篩選")

        c3, c4 = st.columns(2)

        with c3:
            s["selected_source_categories"] = st.multiselect(
                "來源分類",
                options=source_categories,
                default=s.get("selected_source_categories", []),
                key=f"src_cat_{selected_idx}",
            )

            filtered_source_name_options = []
            for item in all_sources:
                item_name = (item.get("name") or "").strip()
                item_cats = item.get("category", []) or []
                if not s["selected_source_categories"]:
                    if item_name and item_name not in filtered_source_name_options:
                        filtered_source_name_options.append(item_name)
                else:
                    if any(cat in s["selected_source_categories"] for cat in item_cats):
                        if item_name and item_name not in filtered_source_name_options:
                            filtered_source_name_options.append(item_name)

            existing_source_names = [
                x for x in s.get("selected_source_names", [])
                if x in filtered_source_name_options
            ]

            s["selected_source_names"] = st.multiselect(
                "來源名稱",
                options=filtered_source_name_options,
                default=existing_source_names,
                key=f"src_names_{selected_idx}",
            )

        with c4:
            s["selected_expert_categories"] = st.multiselect(
                "專家分類",
                options=expert_categories,
                default=s.get("selected_expert_categories", []),
                key=f"exp_cat_{selected_idx}",
            )

            filtered_expert_name_options = []
            for item in all_experts:
                item_name = (item.get("name") or "").strip()
                item_cats = item.get("category", []) or []
                if not s["selected_expert_categories"]:
                    if item_name and item_name not in filtered_expert_name_options:
                        filtered_expert_name_options.append(item_name)
                else:
                    if any(cat in s["selected_expert_categories"] for cat in item_cats):
                        if item_name and item_name not in filtered_expert_name_options:
                            filtered_expert_name_options.append(item_name)

            existing_expert_names = [
                x for x in s.get("selected_expert_names", [])
                if x in filtered_expert_name_options
            ]

            s["selected_expert_names"] = st.multiselect(
                "專家名稱",
                options=filtered_expert_name_options,
                default=existing_expert_names,
                key=f"exp_names_{selected_idx}",
            )

        config["schedules"][selected_idx] = s

        save_c1, save_c2 = st.columns(2)

        with save_c1:
            if st.button("儲存此排程", key=f"save_this_{selected_idx}", use_container_width=True):
                save_auto_export(config)
                st.success("已儲存此排程")

        with save_c2:
            if st.button("刪除此排程", key=f"delete_this_{selected_idx}", use_container_width=True):
                config["schedules"].pop(selected_idx)
                save_auto_export(config)
                st.session_state.automation_selected_index = max(0, selected_idx - 1)
                st.rerun()

    # -------------------------
    # Next Runs Preview
    # -------------------------
    st.markdown("### Next Runs Preview")

    preview_rows = []
    for s in config.get("schedules", []):
        runs = compute_next_runs(s, 5)
        preview_rows.append({
            "name": s.get("name", ""),
            "mode": s.get("schedule_mode", ""),
            "next_5_runs": "\n".join([dt.strftime("%Y-%m-%d %H:%M") for dt in runs]) if runs else "-",
        })

    if preview_rows:
        st.dataframe(pd.DataFrame(preview_rows), use_container_width=True)
    else:
        st.info("目前沒有排程可預覽。")
