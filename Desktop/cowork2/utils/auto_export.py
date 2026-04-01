import inspect
import json
from copy import deepcopy
from pathlib import Path
from datetime import datetime, timedelta

try:
    import pytz
    _TW_TZ = pytz.timezone("Asia/Taipei")
    def _now_tw():
        return datetime.now(_TW_TZ).replace(tzinfo=None)
except Exception:
    def _now_tw():
        return datetime.now()

BASE_DIR = Path(__file__).resolve().parent.parent
CONFIG_DIR = BASE_DIR / "config"
OUTPUT_DIR = BASE_DIR / "output"

AUTO_EXPORT_CONFIG_PATH = CONFIG_DIR / "auto_export.json"
AUTO_EXPORT_STATE_PATH = CONFIG_DIR / "auto_export_state.json"
SOURCES_PATH = CONFIG_DIR / "sources.json"
EXPERTS_PATH = CONFIG_DIR / "experts.json"
PROFILES_PATH = CONFIG_DIR / "profiles.json"
INSIGHTS_PATH = CONFIG_DIR / "insights.txt"

WEEKDAY_LABELS = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]


# =========================================================
# 檔案讀寫
# =========================================================
def safe_load_json(path: Path, default):
    try:
        if path.exists():
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return deepcopy(default)


def safe_save_json(path: Path, data):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def safe_load_text(path: Path, default=""):
    try:
        if path.exists():
            return path.read_text(encoding="utf-8")
    except Exception:
        pass
    return default


# =========================================================
# 舊版相容 / 結構正規化
# =========================================================
def default_schedule():
    return {
        "name": "Morning Briefing",
        "schedule_mode": "daily",
        "once_datetime": "",
        "hourly_interval_hours": 4,
        "daily_times": ["09:00"],
        "weekly_days": [0],
        "weekly_times": ["09:00"],
        "monthly_days": [1],
        "monthly_times": ["09:00"],
        "start_from": "",
        "coverage_hours": 24,
        "selected_source_categories": [],
        "selected_source_names": [],
        "selected_expert_categories": [],
        "selected_expert_names": [],
        "language": "繁體中文",
        "profile": "default",
        "output_formats": ["docx"],
        "output_targets": ["local"],
        "google_drive_folder_id": "",
        "report_mode": "single",
        "multiphase_groups": [],
    }


def normalize_schedule(raw):
    s = default_schedule()
    if isinstance(raw, dict):
        s.update(raw)

    # 舊欄位相容
    if "interval_hours" in s and not s.get("hourly_interval_hours"):
        s["hourly_interval_hours"] = s["interval_hours"]

    if "time" in s and not s.get("daily_times"):
        s["daily_times"] = [s["time"]]

    # 正規化 list 欄位
    for key in [
        "daily_times",
        "weekly_days",
        "weekly_times",
        "monthly_days",
        "monthly_times",
        "selected_source_categories",
        "selected_source_names",
        "selected_expert_categories",
        "selected_expert_names",
        "output_formats",
        "output_targets",
    ]:
        value = s.get(key)
        if value is None:
            s[key] = []
        elif not isinstance(value, list):
            s[key] = [value]

    if not s["daily_times"]:
        s["daily_times"] = ["09:00"]
    if not s["weekly_days"]:
        s["weekly_days"] = [0]
    if not s["weekly_times"]:
        s["weekly_times"] = ["09:00"]
    if not s["monthly_days"]:
        s["monthly_days"] = [1]
    if not s["monthly_times"]:
        s["monthly_times"] = ["09:00"]
    if not s["output_formats"]:
        s["output_formats"] = ["docx"]
    if not s["output_targets"]:
        s["output_targets"] = ["local"]
    if not s.get("profile"):
        s["profile"] = "default"
    if not s.get("language"):
        s["language"] = "繁體中文"
    if not s.get("name"):
        s["name"] = "Untitled Schedule"

    try:
        s["coverage_hours"] = int(s.get("coverage_hours", 24))
    except Exception:
        s["coverage_hours"] = 24

    try:
        s["hourly_interval_hours"] = int(s.get("hourly_interval_hours", 4))
    except Exception:
        s["hourly_interval_hours"] = 4

    s["schedule_mode"] = s.get("schedule_mode", "daily")
    if s["schedule_mode"] not in {"once", "hourly", "daily", "weekly", "monthly"}:
        s["schedule_mode"] = "daily"

    return s


def load_auto_export_config():
    raw = safe_load_json(AUTO_EXPORT_CONFIG_PATH, {"enabled": True, "schedules": []})

    # 舊單一排程格式自動轉換
    if isinstance(raw, dict) and "schedules" not in raw:
        migrated = {
            "enabled": bool(raw.get("enabled", True)),
            "schedules": [normalize_schedule(raw)],
        }
        safe_save_json(AUTO_EXPORT_CONFIG_PATH, migrated)
        return migrated

    config = {
        "enabled": bool(raw.get("enabled", True)),
        "schedules": [],
    }

    schedules = raw.get("schedules", [])
    if isinstance(schedules, list):
        config["schedules"] = [normalize_schedule(x) for x in schedules]

    return config


def save_auto_export_config(config):
    data = {
        "enabled": bool(config.get("enabled", True)),
        "schedules": [normalize_schedule(x) for x in config.get("schedules", [])],
    }
    safe_save_json(AUTO_EXPORT_CONFIG_PATH, data)


def load_auto_export_state():
    return safe_load_json(AUTO_EXPORT_STATE_PATH, {"last_run_keys": {}})


def save_auto_export_state(state):
    safe_save_json(AUTO_EXPORT_STATE_PATH, state)


# =========================================================
# 排程摘要 / editor
# =========================================================
def build_time_interval_text(schedule):
    s = normalize_schedule(schedule)
    mode = s["schedule_mode"]

    if mode == "once":
        return s.get("once_datetime", "")
    if mode == "hourly":
        return f'{s.get("hourly_interval_hours", 4)}h'
    if mode == "daily":
        return ",".join(s.get("daily_times", []))
    if mode == "weekly":
        days = [WEEKDAY_LABELS[d] for d in s.get("weekly_days", []) if 0 <= d <= 6]
        times = s.get("weekly_times", [])
        return f'{"|".join(days)} @ {",".join(times)}'
    if mode == "monthly":
        days = [str(x) for x in s.get("monthly_days", [])]
        times = s.get("monthly_times", [])
        return f'Days {",".join(days)} @ {",".join(times)}'
    return ""


def schedule_to_editor_row(schedule):
    s = normalize_schedule(schedule)
    return {
        "name": s["name"],
        "schedule_mode": s["schedule_mode"],
        "time_or_interval": build_time_interval_text(s),
        "language": s["language"],
        "profile": s["profile"],
        "output_formats": ",".join(s["output_formats"]),
        "output_targets": ",".join(s["output_targets"]),
        "delete": False,
    }


def apply_editor_row_to_schedule(schedule, row):
    s = normalize_schedule(schedule)
    s["name"] = row.get("name", s["name"])
    s["schedule_mode"] = row.get("schedule_mode", s["schedule_mode"])
    s["language"] = row.get("language", s["language"])
    s["profile"] = row.get("profile", s["profile"])

    output_formats = row.get("output_formats", ",".join(s["output_formats"]))
    output_targets = row.get("output_targets", ",".join(s["output_targets"]))
    s["output_formats"] = [x.strip() for x in str(output_formats).split(",") if x.strip()]
    s["output_targets"] = [x.strip() for x in str(output_targets).split(",") if x.strip()]

    time_text = str(row.get("time_or_interval", "")).strip()
    mode = s["schedule_mode"]

    if mode == "hourly":
        txt = time_text.lower().replace("hours", "").replace("hour", "").replace("h", "").strip()
        try:
            s["hourly_interval_hours"] = max(1, int(txt))
        except Exception:
            pass

    elif mode == "daily":
        if time_text:
            s["daily_times"] = [x.strip() for x in time_text.split(",") if x.strip()]

    elif mode == "once":
        if time_text:
            s["once_datetime"] = time_text

    return normalize_schedule(s)


# =========================================================
# 時間工具
# =========================================================
def parse_time_hhmm(text):
    h, m = text.split(":")
    return int(h), int(m)


def safe_parse_datetime(text):
    for fmt in ("%Y-%m-%d %H:%M", "%Y/%m/%d %H:%M", "%Y-%m-%dT%H:%M"):
        try:
            return datetime.strptime(text, fmt)
        except Exception:
            continue
    return None


def candidate_run_times_for_day(schedule, day_dt):
    s = normalize_schedule(schedule)
    mode = s["schedule_mode"]
    out = []

    if mode == "daily":
        for t in s["daily_times"]:
            try:
                h, m = parse_time_hhmm(t)
                out.append(day_dt.replace(hour=h, minute=m, second=0, microsecond=0))
            except Exception:
                pass

    elif mode == "weekly":
        if day_dt.weekday() in s["weekly_days"]:
            for t in s["weekly_times"]:
                try:
                    h, m = parse_time_hhmm(t)
                    out.append(day_dt.replace(hour=h, minute=m, second=0, microsecond=0))
                except Exception:
                    pass

    elif mode == "monthly":
        if day_dt.day in s["monthly_days"]:
            for t in s["monthly_times"]:
                try:
                    h, m = parse_time_hhmm(t)
                    out.append(day_dt.replace(hour=h, minute=m, second=0, microsecond=0))
                except Exception:
                    pass

    return out


def compute_next_runs(schedule, count=5, now=None):
    s = normalize_schedule(schedule)
    now = now or datetime.now()
    mode = s["schedule_mode"]
    results = []

    if mode == "once":
        dt = safe_parse_datetime(s.get("once_datetime", ""))
        if dt and dt >= now:
            return [dt]
        return []

    # 非 once 模式：若設了 start_from，從該日期時間開始算（不早於 now）
    start_from_str = s.get("start_from", "")
    effective_now = now
    if start_from_str:
        try:
            if len(start_from_str) >= 16:
                sf = datetime.strptime(start_from_str[:16], "%Y-%m-%d %H:%M")
            else:
                sf = datetime.strptime(start_from_str[:10], "%Y-%m-%d")
            if sf > now:
                effective_now = sf
        except Exception:
            pass

    if mode == "hourly":
        interval = max(1, int(s.get("hourly_interval_hours", 4)))
        # 對齊整點
        base = effective_now.replace(minute=0, second=0, microsecond=0)
        if effective_now.minute > 0 or effective_now.second > 0 or effective_now.microsecond > 0:
            base += timedelta(hours=1)

        while len(results) < count:
            diff_hours = base.hour % interval
            if diff_hours != 0:
                base += timedelta(hours=(interval - diff_hours))
            results.append(base)
            base += timedelta(hours=interval)
        return results

    # daily / weekly / monthly
    cursor = effective_now.replace(second=0, microsecond=0)
    for i in range(0, 400):
        day_dt = (cursor + timedelta(days=i)).replace(hour=0, minute=0)
        cands = candidate_run_times_for_day(s, day_dt)
        cands = [x for x in cands if x >= effective_now]
        cands.sort()
        for c in cands:
            results.append(c)
            if len(results) >= count:
                return results
    return results


def get_due_run_key(schedule, now=None):
    s = normalize_schedule(schedule)
    now = now or datetime.now()
    mode = s["schedule_mode"]

    if mode == "once":
        dt = safe_parse_datetime(s.get("once_datetime", ""))
        if not dt:
            return None
        if now.year == dt.year and now.month == dt.month and now.day == dt.day and now.hour == dt.hour and now.minute == dt.minute:
            return dt.strftime("%Y-%m-%d %H:%M")
        return None

    # 非 once 模式：若設了 start_from，現在必須 >= 該日期時間才執行
    start_from_str = s.get("start_from", "")
    if start_from_str:
        try:
            if len(start_from_str) >= 16:
                start_from_dt = datetime.strptime(start_from_str[:16], "%Y-%m-%d %H:%M")
            else:
                start_from_dt = datetime.strptime(start_from_str[:10], "%Y-%m-%d")
            if now < start_from_dt:
                return None
        except Exception:
            pass

    if mode == "hourly":
        interval = max(1, int(s.get("hourly_interval_hours", 4)))
        if now.minute != 0:
            return None
        if now.hour % interval != 0:
            return None
        return now.strftime("%Y-%m-%d %H:%M")

    if mode == "daily":
        hhmm = now.strftime("%H:%M")
        if hhmm in s["daily_times"]:
            return now.strftime("%Y-%m-%d %H:%M")
        return None

    if mode == "weekly":
        hhmm = now.strftime("%H:%M")
        if now.weekday() in s["weekly_days"] and hhmm in s["weekly_times"]:
            return now.strftime("%Y-%m-%d %H:%M")
        return None

    if mode == "monthly":
        hhmm = now.strftime("%H:%M")
        if now.day in s["monthly_days"] and hhmm in s["monthly_times"]:
            return now.strftime("%Y-%m-%d %H:%M")
        return None

    return None


def should_run(schedule, state=None, now=None):
    now = now or datetime.now()
    s = normalize_schedule(schedule)
    state = state or load_auto_export_state()
    run_key = get_due_run_key(s, now=now)
    if not run_key:
        return False, None

    last_run_keys = state.setdefault("last_run_keys", {})
    schedule_name = s.get("name", "unnamed")
    if last_run_keys.get(schedule_name) == run_key:
        return False, run_key

    return True, run_key


# =========================================================
# 資料讀取
# =========================================================
def normalize_items(raw):
    if isinstance(raw, list):
        return raw
    if isinstance(raw, dict):
        if "items" in raw and isinstance(raw["items"], list):
            return raw["items"]
        result = []
        for k, v in raw.items():
            if isinstance(v, dict):
                item = deepcopy(v)
                item.setdefault("name", k)
                result.append(item)
        return result
    return []


def filter_items_by_schedule(items, selected_categories, selected_names):
    items = normalize_items(items)

    if not selected_categories and not selected_names:
        return items

    output = []
    for item in items:
        cat_ok = (not selected_categories) or (item.get("category", "未分類") in selected_categories)
        name_ok = (not selected_names) or (item.get("name", "") in selected_names)

        # 只要 category / name 有任一條件組合成立即可保留
        if selected_categories and selected_names:
            if cat_ok and name_ok:
                output.append(item)
        elif selected_categories:
            if cat_ok:
                output.append(item)
        elif selected_names:
            if name_ok:
                output.append(item)
    return output


# =========================================================
# 匯出工具
# =========================================================
def export_text_to_docx(text, output_path: Path, title="公情綜整簡報",
                        start_time=None, end_time=None):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        raise RuntimeError(f"python-docx 未安裝或不可用：{e}")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 標題
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 時間區間（若有提供）
    if start_time and end_time:
        fmt = "%Y-%m-%d %H:%M"
        time_range = f"{start_time.strftime(fmt)} ～ {end_time.strftime(fmt)}"
        p = doc.add_paragraph(time_range)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_paragraph("")  # 空行隔開

    for para in text.split("\n"):
        doc.add_paragraph(para)

    doc.save(str(output_path))
    return str(output_path)


def export_text_to_pdf(text, output_path: Path, title="Briefing Report"):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfgen import canvas
    except Exception as e:
        raise RuntimeError(f"reportlab 未安裝或不可用：{e}")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4

    # 嘗試註冊常見中文字體；若失敗則仍可用 Helvetica
    font_name = "Helvetica"
    candidates = [
        ("NotoSansCJK", BASE_DIR / "assets" / "NotoSansCJKtc-Regular.otf"),
        ("MicrosoftJhengHei", Path("C:/Windows/Fonts/msjh.ttc")),
        ("SimSun", Path("C:/Windows/Fonts/simsun.ttc")),
    ]
    for name, font_path in candidates:
        try:
            if font_path.exists():
                pdfmetrics.registerFont(TTFont(name, str(font_path)))
                font_name = name
                break
        except Exception:
            continue

    y = height - 50
    c.setFont(font_name, 16)
    c.drawString(40, y, title)
    y -= 30

    c.setFont(font_name, 10)
    lines = []
    for para in text.split("\n"):
        if not para.strip():
            lines.append("")
            continue
        chunk = para
        while len(chunk) > 70:
            lines.append(chunk[:70])
            chunk = chunk[70:]
        lines.append(chunk)

    for line in lines:
        if y < 50:
            c.showPage()
            c.setFont(font_name, 10)
            y = height - 50
        c.drawString(40, y, line)
        y -= 14

    c.save()
    return str(output_path)


def upload_to_google_drive_if_needed(file_path, schedule):
    """
    嘗試上傳檔案到 Google Drive。
    回傳 (result_or_None, error_message_or_None)。
    """
    targets = schedule.get("output_targets", [])
    # 同時相容 "google_drive" 和舊版 "drive"
    if "google_drive" not in targets and "drive" not in targets:
        return None, None

    folder_id = schedule.get("google_drive_folder_id", "")
    try:
        from utils.google_drive import upload_to_drive as upload_file_to_drive
    except Exception as e:
        return None, f"google_drive 模組載入失敗：{e}"

    try:
        file_name = Path(file_path).name
        result = upload_file_to_drive(file_path, file_name, folder_id or None)
        return result, None
    except Exception as e:
        return None, str(e)


# =========================================================
# 報告生成主流程
# =========================================================
def try_generate_report_via_report_engine(schedule, context):
    """
    盡量相容既有 report_engine.generate_report 的不同參數形式。
    支援 report_mode: "single" | "multi_phase" | "segmented"。
    """
    import report_engine

    _report_mode = schedule.get("report_mode", "single")

    # ── 分段報告模式 ─────────────────────────────────────────────
    if _report_mode == "segmented":
        if not hasattr(report_engine, "generate_segmented_report"):
            raise RuntimeError("找不到 report_engine.generate_segmented_report")
        report_text = report_engine.generate_segmented_report(
            start_time=context["start_time"],
            end_time=context["end_time"],
            language=schedule.get("language", "繁體中文"),
            insights_text=context.get("insights_text", ""),
            format_options=None,
            status_callback=None,
        )
        if isinstance(report_text, dict):
            report_text = json.dumps(report_text, ensure_ascii=False, indent=2)
        elif not isinstance(report_text, str):
            report_text = str(report_text)
        return report_text

    # ── 單份 / 綜合報告模式 ──────────────────────────────────────
    if not hasattr(report_engine, "generate_report"):
        raise RuntimeError("找不到 report_engine.generate_report")

    fn = report_engine.generate_report
    sig = inspect.signature(fn)

    # multiphase_groups: None = single mode, list = multi-phase (empty list = all groups)
    _multiphase_groups = (
        schedule.get("multiphase_groups") or []
    ) if _report_mode == "multi_phase" else None

    arg_candidates = {
        "schedule": schedule,
        "config": schedule,
        "schedule_config": schedule,
        "coverage_hours": schedule.get("coverage_hours", 24),
        "language": schedule.get("language", "繁體中文"),
        "profile": schedule.get("profile", "default"),
        "selected_source_categories": schedule.get("selected_source_categories", []),
        "selected_source_names": schedule.get("selected_source_names", []),
        "selected_expert_categories": schedule.get("selected_expert_categories", []),
        "selected_expert_names": schedule.get("selected_expert_names", []),
        # 同時用 sources / selected_sources 兩種命名，相容不同版本的 generate_report
        "sources": context["filtered_sources"],
        "selected_sources": context["filtered_sources"] if context["filtered_sources"] else context["all_sources"],
        "experts": context["filtered_experts"],
        "selected_experts": context["filtered_experts"] if context["filtered_experts"] else context["all_experts"],
        "all_sources": context["all_sources"],
        "all_experts": context["all_experts"],
        "profiles": context["profiles"],
        "insights_text": context["insights_text"],
        "insights": context["insights_text"],
        "start_time": context["start_time"],
        "end_time": context["end_time"],
        "multiphase_groups": _multiphase_groups,
    }

    kwargs = {}
    for name in sig.parameters.keys():
        if name in arg_candidates:
            kwargs[name] = arg_candidates[name]

    result = fn(**kwargs)

    # 常見回傳格式兼容
    if isinstance(result, tuple):
        if len(result) >= 1:
            report_text = result[0]
        else:
            report_text = ""
    else:
        report_text = result

    if isinstance(report_text, dict):
        report_text = json.dumps(report_text, ensure_ascii=False, indent=2)
    elif not isinstance(report_text, str):
        report_text = str(report_text)

    return report_text


def run_schedule_job(schedule):
    try:
        schedule = normalize_schedule(schedule)

        all_sources = safe_load_json(SOURCES_PATH, [])
        all_experts = safe_load_json(EXPERTS_PATH, [])
        profiles = safe_load_json(PROFILES_PATH, {"default": {}})
        insights_text = safe_load_text(INSIGHTS_PATH, "")

        filtered_sources = filter_items_by_schedule(
            all_sources,
            schedule.get("selected_source_categories", []),
            schedule.get("selected_source_names", []),
        )
        filtered_experts = filter_items_by_schedule(
            all_experts,
            schedule.get("selected_expert_categories", []),
            schedule.get("selected_expert_names", []),
        )

        end_time = _now_tw()
        start_time = end_time - timedelta(hours=int(schedule.get("coverage_hours", 24)))

        context = {
            "all_sources": normalize_items(all_sources),
            "all_experts": normalize_items(all_experts),
            "filtered_sources": filtered_sources,
            "filtered_experts": filtered_experts,
            "profiles": profiles,
            "insights_text": insights_text,
            "start_time": start_time,
            "end_time": end_time,
        }

        report_text = try_generate_report_via_report_engine(schedule, context)

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        timestamp = _now_tw().strftime("%Y%m%d_%H%M%S")
        base_name = f"公情綜整簡報 {timestamp}"
        files = []

        drive_results = []
        drive_errors = []

        if "docx" in schedule.get("output_formats", []):
            docx_path = OUTPUT_DIR / f"{base_name}.docx"
            files.append(export_text_to_docx(
                report_text, docx_path,
                title="公情綜整簡報",
                start_time=start_time,
                end_time=end_time,
            ))
            dr, de = upload_to_google_drive_if_needed(str(docx_path), schedule)
            if dr:
                drive_results.append(dr.get("webViewLink", "已上傳"))
            if de:
                drive_errors.append(f"docx：{de}")

        if "pdf" in schedule.get("output_formats", []):
            pdf_path = OUTPUT_DIR / f"{base_name}.pdf"
            files.append(export_text_to_pdf(report_text, pdf_path, title=schedule["name"]))
            dr, de = upload_to_google_drive_if_needed(str(pdf_path), schedule)
            if dr:
                drive_results.append(dr.get("webViewLink", "已上傳"))
            if de:
                drive_errors.append(f"pdf：{de}")

        msg_parts = [f"完成：{schedule['name']}"]
        if drive_results:
            msg_parts.append(f"Drive 上傳成功：{len(drive_results)} 個檔案")
        if drive_errors:
            msg_parts.append(f"Drive 上傳失敗：{'; '.join(drive_errors)}")

        return {
            "ok": True,
            "files": files,
            "message": "｜".join(msg_parts),
            "drive_links": drive_results,
            "drive_errors": drive_errors,
            "report_text": report_text,
        }

    except Exception as e:
        return {
            "ok": False,
            "error": str(e),
        }