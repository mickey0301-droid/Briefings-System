import inspect
from datetime import datetime, timedelta, time, date
from pathlib import Path
import pytz

import pandas as pd
import streamlit as st

TW_TZ = pytz.timezone("Asia/Taipei")

def now_tw() -> datetime:
    """回傳台灣時間（naive datetime，去除 tzinfo 以便與現有程式碼相容）"""
    return datetime.now(TW_TZ).replace(tzinfo=None)

# ── 背景排程器（每 60 秒檢查一次到期排程） ──────────────────────────────
# 必須在其他 import 之前啟動，且使用 st.cache_resource 確保只啟動一次
from utils.scheduler_daemon import start_background_scheduler
start_background_scheduler()
# ─────────────────────────────────────────────────────────────────────────

from utils.loaders import (
    load_sources,
    save_sources,
    load_experts,
    save_experts,
    load_profiles,
    save_profiles,
    load_formats,
    save_formats,
    load_insights,
    save_insights,
    load_auto_export,
    save_auto_export,
    load_auto_export_state,
    get_source_categories,
    get_expert_categories,
    source_to_editor_row,
    editor_row_to_source,
    expert_to_editor_row,
    editor_row_to_expert,
    display_expert_name,
    load_category_keywords,
    save_category_keywords,
    DEFAULT_CATEGORY_KEYWORDS,
)
import os
import google.generativeai as genai

# Configure Gemini API before importing report_engine
if "GOOGLE_API_KEY" in st.secrets:
    os.environ["GOOGLE_API_KEY"] = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

import report_engine

try:
    from utils import google_drive as google_drive_utils
except Exception:
    google_drive_utils = None


APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

st.set_page_config(page_title="公情綜整報告", layout="wide")

# 禁止手機瀏覽器的下拉重新整理（pull-to-refresh）
st.markdown(
    """
    <style>
    html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"] {
        overscroll-behavior-y: contain;
        overscroll-behavior: contain;
    }
    </style>
    <script>
    // 攔截 touchstart → touchmove，當頁面已在頂部時阻止原生下拉重整
    (function() {
        var startY = 0;
        document.addEventListener('touchstart', function(e) {
            startY = e.touches[0].clientY;
        }, { passive: true });
        document.addEventListener('touchmove', function(e) {
            var el = e.target;
            // 向上捲動（往下拉）且已在頁面頂部
            if (e.touches[0].clientY > startY && window.scrollY === 0) {
                e.preventDefault();
            }
        }, { passive: false });
    })();
    </script>
    """,
    unsafe_allow_html=True,
)


# =========================================================
# Helpers
# =========================================================
def _now_str():
    return now_tw().strftime("%Y%m%d_%H%M%S")


def _clean_batch_df(df: pd.DataFrame):
    if df is None or df.empty:
        return []
    rows = []
    for _, row in df.iterrows():
        item = {}
        has_value = False
        for col in df.columns:
            val = row[col]
            if pd.isna(val):
                val = ""
            if isinstance(val, str):
                val = val.strip()
            if val not in ["", None]:
                has_value = True
            item[col] = val
        if has_value:
            rows.append(item)
    return rows


def _profiles_map(profiles):
    result = {}
    for p in profiles:
        if isinstance(p, dict):
            name = p.get("name") or p.get("id") or p.get("title") or "Default"
            result[name] = p
        else:
            result[str(p)] = p
    return result


def _filter_sources_by_selection(sources, selected_categories, selected_names):
    filtered = []
    for s in sources:
        if not s.get("enabled", True):
            continue
        cats = s.get("category", []) or []

        # 中共官媒是「主動選取」（opt-in）：只有在使用者明確選了「中共官媒」類別
        # 或明確點選了該來源名稱時，才納入抓取範圍。
        if s.get("type") == "cn_official":
            by_cat  = bool(selected_categories and any(c in selected_categories for c in cats))
            by_name = bool(selected_names and s.get("name") in selected_names)
            if not by_cat and not by_name:
                continue
        else:
            if selected_categories and not any(c in selected_categories for c in cats):
                continue
            if selected_names and s.get("name") not in selected_names:
                continue

        filtered.append(s)
    return filtered


def _filter_experts_by_selection(experts, selected_categories, selected_names):
    filtered = []
    for e in experts:
        if not e.get("enabled", True):
            continue
        cats = e.get("category", []) or []
        name = display_expert_name(e)
        if selected_categories and not any(c in selected_categories for c in cats):
            continue
        if selected_names and name not in selected_names:
            continue
        filtered.append(e)
    return filtered


def _call_generate_report(
    start_time,
    end_time,
    selected_sources,
    selected_experts,
    profile_name,
    language,
    insights_text,
    status_callback=None,
    multiphase_groups=None,
):
    fn = getattr(report_engine, "generate_report", None)
    if fn is None:
        raise RuntimeError("找不到 report_engine.generate_report()")

    sig = inspect.signature(fn)
    kwargs = {}

    candidate_values = {
        "start_time": start_time,
        "end_time": end_time,
        "selected_sources": selected_sources,
        "sources": selected_sources,
        "selected_experts": selected_experts,
        "experts": selected_experts,
        "profile_name": profile_name,
        "profile": profile_name,
        "report_template": profile_name,
        "language": language,
        "insights_text": insights_text,
        "insights": insights_text,
        "status_callback": status_callback,
        "multiphase_groups": multiphase_groups,
    }

    for name in sig.parameters:
        if name in candidate_values:
            kwargs[name] = candidate_values[name]

    result = fn(**kwargs)

    if isinstance(result, tuple):
        if len(result) >= 2:
            return result[0], result[1]
        if len(result) == 1:
            return result[0], []
    return result, []


def _fallback_save_docx(report_text: str, output_path: Path, format_config=None):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title_cfg = (format_config or {}).get("title", {})
    body_cfg = (format_config or {}).get("body", {})
    section_cfg = (format_config or {}).get("section_heading", {})

    title_size = int(title_cfg.get("font_size", 16))
    title_bold = bool(title_cfg.get("bold", True))
    title_align = title_cfg.get("align", "center")

    body_size = int(body_cfg.get("font_size", 12))
    line_spacing = float(body_cfg.get("line_spacing", 1.15))

    section_size = int(section_cfg.get("font_size", 14))
    section_bold = bool(section_cfg.get("bold", True))

    lines = (report_text or "").splitlines()

    for idx, line in enumerate(lines):
        text = line.strip()

        if idx == 0 and text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = title_bold
            run.font.size = Pt(title_size)
            if title_align == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif title_align == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if text.startswith("## ") or text.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(text.replace("## ", "").replace("### ", ""))
            run.bold = section_bold
            run.font.size = Pt(section_size)
            p.paragraph_format.line_spacing = line_spacing
        else:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(body_size)
            p.paragraph_format.line_spacing = line_spacing

    doc.save(str(output_path))
    return str(output_path)


def _call_save_report_docx(report_text: str, items: list, output_path: Path, format_config=None):
    fn = getattr(report_engine, "save_report_docx", None)
    if fn is None:
        return _fallback_save_docx(report_text, output_path, format_config=format_config)

    sig = inspect.signature(fn)
    params = list(sig.parameters.keys())

    try:
        if len(params) == 2:
            return fn(report_text, str(output_path))
        if len(params) == 3:
            return fn(report_text, items, str(output_path))
        if len(params) >= 4:
            return fn(report_text, items, str(output_path), format_config)
        return fn(report_text, str(output_path))
    except Exception:
        return _fallback_save_docx(report_text, output_path, format_config=format_config)


def _save_pdf_from_docx(docx_path: Path, pdf_path: Path):
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return str(pdf_path), None
    except Exception as e:
        return None, str(e)


def _try_upload_to_drive(file_path: str, folder_id: str = ""):
    if google_drive_utils is None:
        return None, "utils/google_drive.py 載入失敗"

    candidate_functions = [
        "upload_file_to_drive",
        "upload_to_drive",
        "upload_file",
    ]

    for fn_name in candidate_functions:
        fn = getattr(google_drive_utils, fn_name, None)
        if fn is None:
            continue
        try:
            sig = inspect.signature(fn)
            kwargs = {}
            if "file_path" in sig.parameters:
                kwargs["file_path"] = file_path
            elif "path" in sig.parameters:
                kwargs["path"] = file_path
            else:
                continue

            if "file_name" in sig.parameters:
                import os
                kwargs["file_name"] = os.path.basename(file_path)

            if "folder_id" in sig.parameters:
                kwargs["folder_id"] = folder_id

            result = fn(**kwargs)
            return result, None
        except Exception as e:
            return None, str(e)

    return None, "找不到可用的 Google Drive upload function"


def _format_output_targets(value):
    if isinstance(value, list):
        return value
    return [value] if value else []


def _format_output_formats(value):
    if isinstance(value, list):
        return value
    return [value] if value else []


def _parse_hhmm(text: str):
    h, m = text.split(":")
    return time(hour=int(h), minute=int(m))


def _next_daily_runs(daily_times, count=5):
    now = now_tw()
    candidates = []
    base_date = now.date()

    for day_offset in range(0, 8):
        d = base_date + timedelta(days=day_offset)
        for t in daily_times:
            try:
                tt = _parse_hhmm(t)
                dt = datetime.combine(d, tt)
                if dt >= now:
                    candidates.append(dt)
            except Exception:
                continue

    candidates.sort()
    return candidates[:count]


def _next_interval_runs(interval_hours, window_start, window_end, count=5):
    now = now_tw()
    interval_hours = max(1, int(interval_hours))
    candidates = []

    try:
        start_t = _parse_hhmm(window_start)
        end_t = _parse_hhmm(window_end)
    except Exception:
        start_t = time(8, 0)
        end_t = time(22, 0)

    for day_offset in range(0, 5):
        d = now.date() + timedelta(days=day_offset)
        day_start = datetime.combine(d, start_t)
        day_end = datetime.combine(d, end_t)

        current = day_start
        while current <= day_end:
            if current >= now:
                candidates.append(current)
            current += timedelta(hours=interval_hours)

    candidates.sort()
    return candidates[:count]


def _append_blank_rows(df: pd.DataFrame, blank_rows: int = 8):
    if df is None:
        df = pd.DataFrame()
    cols = list(df.columns)
    blank = pd.DataFrame([{c: "" for c in cols} for _ in range(blank_rows)])
    if "enabled" in blank.columns:
        blank["enabled"] = True
    if df.empty:
        return blank
    return pd.concat([df, blank], ignore_index=True)


def _build_source_editor_df(source_items, blank_rows=8):
    columns = ["name", "type", "url", "category", "region", "enabled", "description"]
    rows = [source_to_editor_row(x) for x in source_items]
    df = pd.DataFrame(rows, columns=columns) if rows else pd.DataFrame(columns=columns)
    df = _append_blank_rows(df, blank_rows=blank_rows)
    return df


def _build_expert_editor_df(expert_items, blank_rows=8):
    columns = ["name_zh", "name_en", "aliases", "category", "affiliation", "region", "enabled", "description"]
    rows = [expert_to_editor_row(x) for x in expert_items]
    df = pd.DataFrame(rows, columns=columns) if rows else pd.DataFrame(columns=columns)
    df = _append_blank_rows(df, blank_rows=blank_rows)
    return df


# =========================================================
# Load data
# =========================================================
all_sources = load_sources(editable_only=False)
editable_sources = [s for s in all_sources if not s.get("readonly")]
fixed_sources = [s for s in all_sources if s.get("readonly")]

experts = load_experts()
profiles = load_profiles()
formats = load_formats()
profile_map = _profiles_map(profiles)
profile_names = list(profile_map.keys()) if profile_map else ["Strategic Briefing"]
format_names = []
for f in formats:
    if isinstance(f, dict):
        n = f.get("name")
        if n and n not in format_names:
            format_names.append(n)

if "default" not in format_names:
    format_names.insert(0, "default")
saved_insights = load_insights()
auto_export_cfg = load_auto_export()

source_categories = get_source_categories(all_sources)
expert_categories = get_expert_categories(experts)

enabled_source_names = [s["name"] for s in all_sources if s.get("enabled", True)]
enabled_expert_names = [display_expert_name(e) for e in experts if e.get("enabled", True)]


# =========================================================
# Sidebar Navigation
# =========================================================

_NAV_PAGES = [
    ("📋", "Briefings",  "簡報生成"),
    ("💡", "Insights",   "研析方向"),
    ("📰", "Sources",    "來源管理"),
    ("📄", "Formats",    "格式設定"),
    ("⏰", "Schedule",   "排程"),
    ("📊", "Reports",    "報告記錄"),
]

if "selected_page" not in st.session_state:
    st.session_state.selected_page = "Briefings"

with st.sidebar:
    st.markdown(
        "<h1 style='margin-bottom:0.2rem;font-size:1.6rem;font-weight:700;'>公情綜整報告</h1>"
        "<hr style='margin:0.4rem 0 1rem 0;border-color:#e0e0e0;'>",
        unsafe_allow_html=True,
    )
    for _icon, _key, _label in _NAV_PAGES:
        _btn_type = "primary" if st.session_state.selected_page == _key else "secondary"
        if st.button(
            f"{_icon}　{_label}",
            key=f"nav_{_key}",
            use_container_width=True,
            type=_btn_type,
        ):
            st.session_state.selected_page = _key
            st.rerun()

selected_page = st.session_state.selected_page

_PAGE_TITLES = {
    "Briefings": "簡報生成",
    "Insights":  "研析方向",
    "Sources":   "來源管理",
    "Formats":   "格式設定",
    "Schedule":  "排程",
    "Reports":   "報告記錄",
}
st.title(_PAGE_TITLES.get(selected_page, selected_page))


# =========================================================
# Briefings
# =========================================================
if selected_page == "Briefings":

    now = now_tw()
    default_start = now - timedelta(hours=24)
    default_end = now

    c1, c2 = st.columns(2)
    with c1:
        start_date = st.date_input("開始日期", value=default_start.date(), key="brief_start_date")
        start_time_input = st.time_input(
            "開始時間",
            value=default_start.time().replace(second=0, microsecond=0),
            key="brief_start_time"
        )
    with c2:
        end_date = st.date_input("結束日期", value=default_end.date(), key="brief_end_date")
        end_time_input = st.time_input(
            "結束時間",
            value=default_end.time().replace(second=0, microsecond=0),
            key="brief_end_time"
        )

    start_dt = datetime.combine(start_date, start_time_input)
    end_dt = datetime.combine(end_date, end_time_input)

    st.markdown("---")

    c1, c2 = st.columns(2)
    with c1:
        selected_source_categories = st.multiselect(
            "來源分類",
            options=source_categories,
            default=[],
            key="brief_source_categories",
        )
        brief_source_options = [s["name"] for s in all_sources if s.get("enabled", True)]
        selected_source_names = st.multiselect(
            "來源",
            options=brief_source_options,
            default=[],
            key="brief_source_names",
        )

    with c2:
        selected_expert_categories = st.multiselect(
            "專家分類",
            options=expert_categories,
            default=[],
            key="brief_expert_categories",
        )
        selected_expert_names = st.multiselect(
            "專家",
            options=enabled_expert_names,
            default=[],
            key="brief_expert_names",
        )

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        profile_name = st.selectbox(
            "報告模板",
            options=profile_names,
            index=0,
            key="briefings_profile_selectbox",
        )
    with c2:
        selected_format_name = st.selectbox(
            "Format",
            options=format_names,
            index=format_names.index("default") if "default" in format_names else 0,
            key="briefings_format_selectbox",
        )
    with c3:
        language = st.selectbox(
            "語言",
            options=["繁體中文", "英文", "日文", "簡體中文"],
            index=0,
            key="briefings_language_selectbox",
        )
    with c4:
        output_formats = st.multiselect(
            "輸出格式",
            options=["docx", "pdf"],
            default=["docx"],
            key="briefings_output_formats",
        )

    # 讀取預設輸出設定
    default_output_targets = auto_export_cfg.get("default_output_targets", ["local"])
    default_drive_folder = auto_export_cfg.get("default_drive_folder_id", "")
    default_local_folder = auto_export_cfg.get("default_local_folder", "outputs")

    c5, c6 = st.columns(2)
    with c5:
        output_targets = st.multiselect(
            "輸出位置",
            options=["local", "google_drive"],
            default=default_output_targets,
            key="briefings_output_targets",
        )
        local_folder = st.text_input(
            "Local 輸出資料夾",
            value=default_local_folder,
            key="briefings_local_folder",
        )
    with c6:
        # Google Drive 資料夾選擇（同 Schedule 頁面邏輯）
        _df_list = auto_export_cfg.get("drive_folders", [])
        _df_names = [f.get("name", "") or f.get("folder_id", "") for f in _df_list]
        _cur_fid = default_drive_folder
        _cur_idx = next(
            (i for i, f in enumerate(_df_list) if f.get("folder_id") == _cur_fid),
            None,
        )
        if _df_names:
            _options = ["（手動輸入）"] + _df_names
            _sel = st.selectbox(
                "Google Drive 資料夾",
                options=_options,
                index=(_cur_idx + 1) if _cur_idx is not None else 0,
                key="briefings_gdrive_sel",
            )
            if _sel == "（手動輸入）":
                google_drive_folder_id = st.text_input(
                    "Folder ID（手動輸入）",
                    value=_cur_fid,
                    key="briefings_gdrive_manual",
                )
            else:
                _chosen = next(
                    (f for f in _df_list if (f.get("name") or f.get("folder_id")) == _sel),
                    None,
                )
                google_drive_folder_id = _chosen.get("folder_id", "") if _chosen else ""
                st.caption(f"Folder ID：`{google_drive_folder_id}`")
        else:
            google_drive_folder_id = st.text_input(
                "Google Drive Folder ID",
                value=_cur_fid,
                key="briefings_gdrive",
                help="先在 Schedule 頁面的「Google Drive 資料夾」區塊新增資料夾，之後可在此選擇。",
            )

    # ── 報告模式 ────────────────────────────────────────────────────────
    _rmode_col1, _rmode_col2 = st.columns([1, 2])
    with _rmode_col1:
        def _rmode_label(x):
            if x == "single":
                return "單份報告"
            elif x == "multi_phase":
                return "綜合報告（多段生成）"
            else:
                return "分段報告（按章節搜尋）"
        report_mode_brief = st.radio(
            "報告模式",
            options=["single", "multi_phase", "segmented"],
            format_func=_rmode_label,
            horizontal=True,
            key="briefings_report_mode",
        )
    with _rmode_col2:
        if report_mode_brief == "multi_phase":
            _gmap = {k: report_engine.MULTIPHASE_GROUP_OPTIONS and report_engine._MULTIPHASE_GROUP_ZH.get(k, k)
                     for k in report_engine.MULTIPHASE_GROUP_OPTIONS}
            multiphase_groups_brief = st.multiselect(
                "包含來源群組（空白 = 全部）",
                options=list(_gmap.keys()),
                format_func=lambda x: _gmap.get(x, x),
                default=[],
                key="briefings_multiphase_groups",
            )
        elif report_mode_brief == "segmented":
            multiphase_groups_brief = None
            st.caption("分段報告：每章每節獨立搜尋 Google News，生成各節小報告後彙整成完整報告。無需選擇來源群組。")
        else:
            multiphase_groups_brief = None
            st.caption("單份報告：所有來源一次性生成。")

    extra_insights = st.text_area(
        "本次補充 insights",
        value="",
        height=180,
        key="briefings_extra_insights",
    )

    selected_sources = _filter_sources_by_selection(
        all_sources,
        selected_source_categories,
        selected_source_names,
    )
    selected_experts = _filter_experts_by_selection(
        experts,
        selected_expert_categories,
        selected_expert_names,
    )

    s1, s2, s3 = st.columns(3)
    s1.metric("已選來源數", len(selected_sources))
    s2.metric("已選專家數", len(selected_experts))
    s3.metric("時間範圍（小時）", round(max((end_dt - start_dt).total_seconds() / 3600, 0), 2))

    if st.button("一鍵生成並輸出", type="primary", use_container_width=True, key="run_briefings"):
        if start_dt >= end_dt:
            st.error("開始時間必須早於結束時間。")
        elif not output_formats:
            st.error("請至少選一種輸出格式。")
        elif not output_targets:
            st.error("請至少選一種輸出位置。")
        else:
            status = st.empty()
            detail = st.empty()
            progress = st.progress(0)
            result_box = st.container()

            try:
                status.info("開始生成")
                detail.caption("正在準備任務參數。")
                progress.progress(5)

                insight_lines = []

                if isinstance(saved_insights, list):
                    for ins in saved_insights:
                        if not isinstance(ins, dict):
                            continue

                        title = str(ins.get("title", "")).strip()
                        content = str(ins.get("content", "")).strip()

                        if title and content:
                            insight_lines.append(f"{title}: {content}")
                        elif content:
                            insight_lines.append(content)

                elif isinstance(saved_insights, str) and saved_insights.strip():
                    insight_lines.append(saved_insights.strip())

                combined_insights = "\n".join(insight_lines)
                selected_format_config = next(
                    (f for f in formats if f.get("name") == selected_format_name),
                    next((f for f in formats if f.get("name") == "default"), None)
                )
                if extra_insights.strip():
                    combined_insights = f"{combined_insights}\n\n{extra_insights.strip()}".strip()

                status.info("抓取資料中…")

                # ── 即時抓取狀態顯示 ──────────────────────────────────
                fetch_log_placeholder = st.empty()
                fetch_log_lines: list[str] = []
                rss_found_sources: list[str] = []

                def _on_fetch_status(event, detail, *args):
                    if event == "stage":
                        fetch_log_lines.append(detail)
                        if "🤖" in detail:
                            status.info("AI 生成簡報中…")
                            progress.progress(75)
                        elif "✅ 全文" in detail:
                            progress.progress(65)
                        elif "📄" in detail:
                            progress.progress(55)
                        elif "✅ RSS" in detail:
                            progress.progress(30)
                        elif "⏳" in detail:
                            progress.progress(10)
                    elif event == "rss" and detail:
                        rss_found_sources.append(detail)
                        completed, total, total_items = args[0], args[1], args[2]
                        pct = int(10 + (completed / max(total, 1)) * 20)
                        progress.progress(min(pct, 30))
                    # 組合顯示文字
                    display_lines = fetch_log_lines[-6:]
                    if rss_found_sources:
                        recent = "　".join(rss_found_sources[-4:])
                        display_lines.append(f"　　↳ {recent}")
                    fetch_log_placeholder.markdown(
                        "\n\n".join(display_lines)
                    )

                progress.progress(10)

                if report_mode_brief == "segmented":
                    seg_fn = getattr(report_engine, "generate_segmented_report", None)
                    if seg_fn is None:
                        raise RuntimeError("找不到 report_engine.generate_segmented_report()")
                    report_text, filtered_items = seg_fn(
                        start_time=start_dt,
                        end_time=end_dt,
                        language=language,
                        insights_text=combined_insights,
                        format_options=selected_format_config,
                        status_callback=_on_fetch_status,
                    )
                else:
                    report_text, filtered_items = _call_generate_report(
                        start_time=start_dt,
                        end_time=end_dt,
                        selected_sources=selected_sources,
                        selected_experts=selected_experts,
                        profile_name=profile_name,
                        language=language,
                        insights_text=combined_insights,
                        status_callback=_on_fetch_status,
                        multiphase_groups=(
                            multiphase_groups_brief if report_mode_brief == "multi_phase" else None
                        ),
                    )

                fetch_log_placeholder.empty()
                status.info("輸出檔案")
                progress.progress(80)

                ts = _now_str()
                base_name = f"briefings_{ts}"
                output_dir = Path(local_folder)
                output_dir.mkdir(parents=True, exist_ok=True)

                docx_path = output_dir / f"{base_name}.docx"
                pdf_path = output_dir / f"{base_name}.pdf"

                output_logs = []
                generated_files = []

                if "docx" in output_formats or "pdf" in output_formats:
                    _call_save_report_docx(report_text, filtered_items, docx_path, selected_format_config)
                    generated_files.append(docx_path)

                if "pdf" in output_formats:
                    pdf_result, pdf_error = _save_pdf_from_docx(docx_path, pdf_path)
                    if pdf_result:
                        generated_files.append(pdf_path)
                    else:
                        output_logs.append(f"PDF 轉換失敗：{pdf_error}")

                if "local" in output_targets:
                    for f in generated_files:
                        output_logs.append(f"已輸出到本機：{f}")

                if "google_drive" in output_targets:
                    for f in generated_files:
                        uploaded, err = _try_upload_to_drive(str(f), google_drive_folder_id)
                        if err:
                            output_logs.append(f"Google Drive 上傳失敗（{f.name}）：{err}")
                        else:
                            output_logs.append(f"Google Drive 已上傳：{f.name}")

                progress.progress(100)
                status.success("完成")
                # 儲存預設輸出設定
                auto_export_cfg["default_output_targets"] = output_targets
                auto_export_cfg["default_drive_folder_id"] = google_drive_folder_id
                auto_export_cfg["default_local_folder"] = local_folder

                save_auto_export(auto_export_cfg)

                with result_box:
                    st.markdown("### 執行結果")
                    st.write(f"**本次時間範圍：** {start_dt} ～ {end_dt}")
                    st.write(f"**已選來源數：** {len(selected_sources)}")
                    st.write(f"**已選專家數：** {len(selected_experts)}")
                    st.write(f"**已取得文章數：** {len(filtered_items) if filtered_items is not None else 0}")

                    if output_logs:
                        st.markdown("**輸出結果：**")
                        for line in output_logs:
                            st.write(f"- {line}")

                    st.markdown("### 報告預覽")
                    st.text_area("report", value=report_text or "", height=400, label_visibility="collapsed", key="briefings_report_preview")

            except Exception as e:
                status.error(f"生成失敗：{e}")


# =========================================================
# Insights
# =========================================================
elif selected_page == "Insights":

    st.caption("Strategic guidance used by AI when generating reports. These will NOT be cited.")

    insights_data = load_insights()

    # 相容舊版：如果還是 str，就先轉成空列表
    if isinstance(insights_data, str):
        insights_data = []
    elif not isinstance(insights_data, list):
        insights_data = []

    table_rows = []
    for item in insights_data:
        if not isinstance(item, dict):
            continue

        tags_value = item.get("tags", [])
        if isinstance(tags_value, list):
            tags_str = ", ".join([str(x).strip() for x in tags_value if str(x).strip()])
        elif isinstance(tags_value, str):
            tags_str = tags_value
        else:
            tags_str = ""

        table_rows.append({
            "title": item.get("title", ""),
            "content": item.get("content", ""),
            "tags": tags_str,
        })

    df = pd.DataFrame(table_rows, columns=["title", "content", "tags"])

    # 新增空白列讓 editor 可新增
    blank_rows = pd.DataFrame([
        {"title": "", "content": "", "tags": ""}
        for _ in range(10)
    ])

    df = pd.concat([df, blank_rows], ignore_index=True)

    edited_df = st.data_editor(
        df,
        num_rows="dynamic",
        use_container_width=True,
        height=420,
        key="insights_editor",
        column_config={
            "title": st.column_config.TextColumn(
                "Title",
                width="medium",
                help="Short headline for this insight"
            ),
            "content": st.column_config.TextColumn(
                "Content",
                width="large",
                help="Strategic guidance for the AI"
            ),
            "tags": st.column_config.TextColumn(
            "Tags",
            width="medium",
            help="Comma separated topics (optional)"
        ),
    },
)

    if st.button("Save Insights", key="save_insights_btn", use_container_width=True):
        cleaned = []

        for _, row in edited_df.iterrows():
            title = str(row.get("title", "")).strip()
            content = str(row.get("content", "")).strip()
            tags_raw = str(row.get("tags", "")).strip()

            if not title and not content and not tags_raw:
                continue

            tag_list = [t.strip() for t in tags_raw.split(",") if t.strip()]

            cleaned.append({
                "title": title,
                "content": content,
                "tags": tag_list,
            })

        save_insights(cleaned)
        st.success("Insights saved.")
        st.rerun()


# =========================================================
# Sources
# =========================================================
elif selected_page == "Sources":

    # 載入關鍵字設定（每次頁面渲染時讀取，儲存後下次 rerun 即生效）
    _cat_kw = load_category_keywords()

    # ── 版本計數器：每次來源儲存後加一，強制 data_editor 重置為最新資料 ────────
    # 防止 Streamlit data_editor 的 session_state 沿用舊狀態，
    # 導致新增來源在後續「儲存XX媒體編輯」時被覆蓋消失。
    _src_v = st.session_state.get("_src_version", 0)

    src_tab_add, src_tab_tw, src_tab_intl, src_tab_experts, src_tab_global, src_tab_cn = st.tabs([
        "新增來源", "自訂台灣媒體", "自訂國際媒體", "自訂專家", "全球媒體", "中國媒體"
    ])

    tw_sources = [s for s in editable_sources if "自訂台灣媒體" in (s.get("category") or [])]
    intl_sources = [s for s in editable_sources if "自訂國際媒體" in (s.get("category") or [])]
    global_sources_ui = [s for s in all_sources if "全球媒體" in (s.get("category") or [])]

    # ── 新增來源 ──────────────────────────────────────────────────────────────
    with src_tab_add:
        target_cat = st.selectbox(
            "加入至媒體分類",
            options=["自訂台灣媒體", "自訂國際媒體"],
            key="src_add_target_cat",
        )

        with st.expander("單筆新增來源", expanded=False):
            c1, c2 = st.columns(2)
            with c1:
                src_name = st.text_input("name", key="single_src_name")
                src_type = st.selectbox("type", options=["rss", "domain"], key="single_src_type")
                src_url = st.text_input("url", key="single_src_url")
            with c2:
                src_region = st.text_input("region", key="single_src_region")
                src_enabled = st.checkbox("enabled", value=True, key="single_src_enabled")
                src_description = st.text_area("description", key="single_src_description", height=120)

            if st.button("新增來源", key="add_single_source"):
                new_item = editor_row_to_source({
                    "name": src_name,
                    "type": src_type,
                    "url": src_url,
                    "category": target_cat,
                    "region": src_region,
                    "enabled": src_enabled,
                    "description": src_description,
                })
                current = load_sources(editable_only=True)
                if not new_item["name"]:
                    st.error("來源名稱不可空白。")
                else:
                    current.append(new_item)
                    save_sources(current)
                    st.success(f"已新增來源至「{target_cat}」。")
                    st.session_state["_src_version"] = _src_v + 1
                    st.rerun()

        st.markdown("### 批次貼上新增來源")
        st.caption(f"複製多列資料貼到下表，再按「批次加入」，category 將自動設為「{target_cat}」。")

        _batch_cols = ["name", "type", "url", "region", "enabled", "description"]
        _src_batch_default = pd.DataFrame([{c: "" for c in _batch_cols} for _ in range(8)])
        _src_batch_default["enabled"] = True

        source_batch_df = st.data_editor(
            _src_batch_default,
            num_rows="dynamic",
            use_container_width=True,
            height=280,
            key=f"source_batch_editor_{_src_v}",
            column_config={
                "name": st.column_config.TextColumn("name"),
                "type": st.column_config.SelectboxColumn("type", options=["rss", "domain"]),
                "url": st.column_config.TextColumn("url"),
                "region": st.column_config.TextColumn("region"),
                "enabled": st.column_config.CheckboxColumn("enabled", default=True),
                "description": st.column_config.TextColumn("description"),
            },
        )

        if st.button("批次加入來源", key="batch_add_sources"):
            rows = _clean_batch_df(source_batch_df)
            if not rows:
                st.warning("沒有可加入的來源資料。")
            else:
                current = load_sources(editable_only=True)
                name_set = {x.get("name", "").strip() for x in current}
                added = 0
                for row in rows:
                    row["category"] = target_cat
                    item = editor_row_to_source(row)
                    if not item["name"]:
                        continue
                    if item["name"] in name_set:
                        current = [x for x in current if x.get("name") != item["name"]]
                    current.append(item)
                    name_set.add(item["name"])
                    added += 1
                save_sources(current)
                st.success(f"已批次加入 {added} 筆來源至「{target_cat}」。")
                st.session_state["_src_version"] = _src_v + 1
                st.rerun()

    # ── 自訂台灣媒體 ──────────────────────────────────────────────────────────
    with src_tab_tw:
        with st.expander("🔍 Google News RSS 關鍵字篩選", expanded=False):
            st.caption("此類別的 domain 來源會以下列關鍵字向 Google News 查詢，只抓取符合的報導。用 OR 分隔多個關鍵字，留空代表不篩選。")
            _kw_tw = st.text_area(
                "自訂台灣媒體 關鍵字",
                value=_cat_kw.get("自訂台灣媒體", DEFAULT_CATEGORY_KEYWORDS.get("自訂台灣媒體", "")),
                height=80,
                key="kw_editor_tw",
                label_visibility="collapsed",
            )
            if st.button("儲存關鍵字", key="save_kw_tw", use_container_width=True):
                _cat_kw["自訂台灣媒體"] = _kw_tw.strip()
                save_category_keywords(_cat_kw)
                st.success("關鍵字已儲存。")
                st.rerun()
        st.caption(f"共 {len(tw_sources)} 筆")
        tw_df = _build_source_editor_df(tw_sources, blank_rows=5)
        edited_tw_df = st.data_editor(
            tw_df,
            num_rows="dynamic",
            use_container_width=True,
            height=420,
            key=f"tw_sources_editor_{_src_v}",
            column_config={
                "name": st.column_config.TextColumn("name"),
                "type": st.column_config.SelectboxColumn("type", options=["rss", "domain"]),
                "url": st.column_config.TextColumn("url"),
                "category": st.column_config.TextColumn("category"),
                "region": st.column_config.TextColumn("region"),
                "enabled": st.column_config.CheckboxColumn("enabled", default=True),
                "description": st.column_config.TextColumn("description"),
            },
        )
        c1, c2, c3 = st.columns([2, 2, 1])
        with c1:
            if st.button("儲存台灣媒體編輯", key="save_tw_sources", use_container_width=True):
                rows = _clean_batch_df(edited_tw_df)
                current = load_sources(editable_only=True)
                non_tw = [s for s in current if "自訂台灣媒體" not in (s.get("category") or [])]
                new_tw = []
                for row in rows:
                    item = editor_row_to_source(row)
                    if item["name"]:
                        if "自訂台灣媒體" not in (item.get("category") or []):
                            item["category"] = ["自訂台灣媒體"]
                        new_tw.append(item)
                save_sources(non_tw + new_tw)
                st.success("台灣媒體清單已儲存。")
                st.session_state["_src_version"] = _src_v + 1
                st.rerun()
        with c2:
            del_tw = st.multiselect("刪除來源", options=[s["name"] for s in tw_sources], key="delete_tw_names")
            if st.button("刪除選取", key="delete_tw_btn", use_container_width=True):
                current = load_sources(editable_only=True)
                current = [x for x in current if x.get("name") not in del_tw]
                save_sources(current)
                st.success(f"已刪除 {len(del_tw)} 筆。")
                st.session_state["_src_version"] = _src_v + 1
                st.rerun()
        with c3:
            if st.button("🔬 測試抓取", key="test_tw_fetch", use_container_width=True):
                st.session_state["show_tw_test"] = True
        if st.session_state.get("show_tw_test"):
            with st.spinner("測試抓取中（走正式 pipeline）..."):
                import report_engine as _re
                _now = datetime.now()
                _start = _now - timedelta(hours=24)
                _diag = [_re.debug_fetch_source(_src, start_time=_start, end_time=_now)
                         for _src in tw_sources[:3]]
            st.markdown("**診斷結果（與正式報告走完全相同路徑）：**")
            for _d in _diag:
                _icon = "✅" if _d["items_parsed"] > 0 else ("❌" if _d["error"] else "⚠️")
                st.markdown(
                    f"{_icon} **{_d['name']}** &nbsp;·&nbsp; "
                    f"type=`{_d['src_type']}` &nbsp;·&nbsp; "
                    f"HTTP `{_d['http_status']}` &nbsp;·&nbsp; "
                    f"**{_d['items_parsed']} 篇**"
                )
                if _d["error"]:
                    st.error(_d["error"])
                with st.expander("詳細", expanded=_d["items_parsed"] == 0):
                    st.write(f"**原始 url：** `{_d['original_url']}`")
                    st.write(f"**抽出 domain：** `{_d['domain_extracted']}`")
                    st.write(f"**content-type：** `{_d['content_type']}`")
                    st.write(f"**回應長度：** {_d['response_len']} 字元")
                    st.code(_d["rss_url"], language=None)
                    if _d["response_preview"]:
                        st.caption(f"回應前 300 字：{_d['response_preview'][:300]}")
                    for _it in _d["items"][:3]:
                        st.markdown(f"- {_it.get('title','')[:80]}　`{_it.get('published','')[:20]}`")
            st.session_state["show_tw_test"] = False

    # ── 自訂國際媒體 ──────────────────────────────────────────────────────────
    with src_tab_intl:
        with st.expander("🔍 Google News RSS 關鍵字篩選", expanded=False):
            st.caption("此類別的 domain 來源會以下列關鍵字向 Google News 查詢，只抓取符合的報導。用 OR 分隔多個關鍵字，留空代表不篩選。")
            _kw_intl = st.text_area(
                "自訂國際媒體 關鍵字",
                value=_cat_kw.get("自訂國際媒體", DEFAULT_CATEGORY_KEYWORDS.get("自訂國際媒體", "")),
                height=80,
                key="kw_editor_intl",
                label_visibility="collapsed",
            )
            if st.button("儲存關鍵字", key="save_kw_intl", use_container_width=True):
                _cat_kw["自訂國際媒體"] = _kw_intl.strip()
                save_category_keywords(_cat_kw)
                st.success("關鍵字已儲存。")
                st.rerun()
        st.caption(f"共 {len(intl_sources)} 筆")
        intl_df = _build_source_editor_df(intl_sources, blank_rows=5)
        edited_intl_df = st.data_editor(
            intl_df,
            num_rows="dynamic",
            use_container_width=True,
            height=420,
            key=f"intl_sources_editor_{_src_v}",
            column_config={
                "name": st.column_config.TextColumn("name"),
                "type": st.column_config.SelectboxColumn("type", options=["rss", "domain"]),
                "url": st.column_config.TextColumn("url"),
                "category": st.column_config.TextColumn("category"),
                "region": st.column_config.TextColumn("region"),
                "enabled": st.column_config.CheckboxColumn("enabled", default=True),
                "description": st.column_config.TextColumn("description"),
            },
        )
        c1, c2 = st.columns(2)
        with c1:
            if st.button("儲存國際媒體編輯", key="save_intl_sources", use_container_width=True):
                rows = _clean_batch_df(edited_intl_df)
                current = load_sources(editable_only=True)
                non_intl = [s for s in current if "自訂國際媒體" not in (s.get("category") or [])]
                new_intl = []
                for row in rows:
                    item = editor_row_to_source(row)
                    if item["name"]:
                        if "自訂國際媒體" not in (item.get("category") or []):
                            item["category"] = ["自訂國際媒體"]
                        new_intl.append(item)
                save_sources(non_intl + new_intl)
                st.success("國際媒體清單已儲存。")
                st.session_state["_src_version"] = _src_v + 1
                st.rerun()
        with c2:
            del_intl = st.multiselect("刪除來源", options=[s["name"] for s in intl_sources], key="delete_intl_names")
            if st.button("刪除選取", key="delete_intl_btn", use_container_width=True):
                current = load_sources(editable_only=True)
                current = [x for x in current if x.get("name") not in del_intl]
                save_sources(current)
                st.success(f"已刪除 {len(del_intl)} 筆。")
                st.session_state["_src_version"] = _src_v + 1
                st.rerun()

    # ── 自訂專家 ──────────────────────────────────────────────────────────────
    with src_tab_experts:

        with st.expander("🔍 Google News RSS 關鍵字篩選", expanded=False):
            st.caption("此類別的 domain 來源會以下列關鍵字向 Google News 查詢，只抓取符合的報導。用 OR 分隔多個關鍵字，留空代表不篩選。")
            _kw_experts = st.text_area(
                "自訂專家 關鍵字",
                value=_cat_kw.get("自訂專家", DEFAULT_CATEGORY_KEYWORDS.get("自訂專家", "")),
                height=80,
                key="kw_editor_experts",
                label_visibility="collapsed",
            )
            if st.button("儲存關鍵字", key="save_kw_experts", use_container_width=True):
                _cat_kw["自訂專家"] = _kw_experts.strip()
                save_category_keywords(_cat_kw)
                st.success("關鍵字已儲存。")
                st.rerun()

        with st.expander("單筆新增專家", expanded=False):
            c1, c2 = st.columns(2)
            with c1:
                exp_name_zh = st.text_input("中文名 name_zh", key="single_exp_name_zh")
                exp_name_en = st.text_input("英文名 name_en", key="single_exp_name_en")
                exp_aliases = st.text_input("aliases（逗號分隔）", key="single_exp_aliases")
                exp_category = st.text_input("category（逗號分隔）", key="single_exp_category")
            with c2:
                exp_affiliation = st.text_input("affiliation", key="single_exp_affiliation")
                exp_region = st.text_input("region", key="single_exp_region")
                exp_enabled = st.checkbox("enabled", value=True, key="single_exp_enabled")
                exp_description = st.text_area("description", key="single_exp_description", height=120)

            if st.button("新增專家", key="add_single_expert"):
                new_item = editor_row_to_expert(
                    {
                        "name_zh": exp_name_zh,
                        "name_en": exp_name_en,
                        "aliases": exp_aliases,
                        "category": exp_category,
                        "affiliation": exp_affiliation,
                        "region": exp_region,
                        "enabled": exp_enabled,
                        "description": exp_description,
                    }
                )
                current = load_experts()
                display_name = display_expert_name(new_item)
                if not display_name or display_name == "Unnamed Expert":
                    st.error("至少要填中文名或英文名。")
                else:
                    current.append(new_item)
                    save_experts(current)
                    st.success("已新增專家。")
                    st.rerun()

        st.markdown("### 表格式批次貼上新增專家")
        st.caption("可直接從外部複製多列資料貼到下表，再按「批次加入專家」。")

        expert_batch_columns = ["name_zh", "name_en", "aliases", "category", "affiliation", "region", "enabled", "description"]
        expert_batch_default = pd.DataFrame([{c: "" for c in expert_batch_columns} for _ in range(8)])
        expert_batch_default["enabled"] = True

        expert_batch_df = st.data_editor(
            expert_batch_default,
            num_rows="dynamic",
            use_container_width=True,
            height=280,
            key="expert_batch_editor",
            column_config={
                "name_zh": st.column_config.TextColumn("name_zh"),
                "name_en": st.column_config.TextColumn("name_en"),
                "aliases": st.column_config.TextColumn("aliases"),
                "category": st.column_config.TextColumn("category"),
                "affiliation": st.column_config.TextColumn("affiliation"),
                "region": st.column_config.TextColumn("region"),
                "enabled": st.column_config.CheckboxColumn("enabled", default=True),
                "description": st.column_config.TextColumn("description"),
            },
        )

        if st.button("批次加入專家", key="batch_add_experts"):
            rows = _clean_batch_df(expert_batch_df)
            if not rows:
                st.warning("沒有可加入的專家資料。")
            else:
                current = load_experts()
                name_set = {display_expert_name(x) for x in current}
                added = 0
                for row in rows:
                    item = editor_row_to_expert(row)
                    disp = display_expert_name(item)
                    if not disp or disp == "Unnamed Expert":
                        continue
                    if disp in name_set:
                        current = [x for x in current if display_expert_name(x) != disp]
                    current.append(item)
                    name_set.add(disp)
                    added += 1
                save_experts(current)
                st.success(f"已批次加入 / 更新 {added} 筆專家。")
                st.rerun()

        st.markdown("### 既有專家清單")
        st.caption("這裡也可以直接貼上多列資料、修改既有資料、增加新列，再按儲存。")

        experts_df = _build_expert_editor_df(experts, blank_rows=10)

        edited_experts_df = st.data_editor(
            experts_df,
            num_rows="dynamic",
            use_container_width=True,
            height=420,
            key="editable_experts_editor",
            column_config={
                "name_zh": st.column_config.TextColumn("name_zh"),
                "name_en": st.column_config.TextColumn("name_en"),
                "aliases": st.column_config.TextColumn("aliases"),
                "category": st.column_config.TextColumn("category"),
                "affiliation": st.column_config.TextColumn("affiliation"),
                "region": st.column_config.TextColumn("region"),
                "enabled": st.column_config.CheckboxColumn("enabled", default=True),
                "description": st.column_config.TextColumn("description"),
            },
        )

        c1, c2 = st.columns(2)
        with c1:
            if st.button("儲存專家清單編輯", key="save_experts_table", use_container_width=True):
                rows = _clean_batch_df(edited_experts_df)
                cleaned = []
                for row in rows:
                    item = editor_row_to_expert(row)
                    disp = display_expert_name(item)
                    if disp and disp != "Unnamed Expert":
                        cleaned.append(item)
                save_experts(cleaned)
                st.success("專家清單已儲存。")
                st.rerun()

        with c2:
            expert_delete_options = [display_expert_name(x) for x in experts]
            delete_expert_names = st.multiselect("刪除專家", options=expert_delete_options, key="delete_expert_names")
            if st.button("刪除選取專家", key="delete_experts_btn", use_container_width=True):
                current = load_experts()
                current = [x for x in current if display_expert_name(x) not in delete_expert_names]
                save_experts(current)
                st.success(f"已刪除 {len(delete_expert_names)} 筆專家。")
                st.rerun()

        st.markdown("### 專家搜尋名稱預覽")
        preview_rows = []
        for e in load_experts():
            preview_rows.append(
                {
                    "display_name": display_expert_name(e),
                    "search_names": ", ".join(e.get("search_names", [])),
                    "aliases": ", ".join(e.get("aliases", [])),
                    "category": ", ".join(e.get("category", [])),
                    "affiliation": e.get("affiliation", ""),
                    "region": e.get("region", ""),
                    "enabled": e.get("enabled", True),
                }
            )
        if preview_rows:
            st.dataframe(pd.DataFrame(preview_rows), use_container_width=True, hide_index=True)
        else:
            st.info("目前尚無專家資料。")


    # ── 全球媒體 ──────────────────────────────────────────────────────────────
    with src_tab_global:
        with st.expander("🔍 Google News RSS 關鍵字篩選", expanded=False):
            st.caption("此類別的 domain 來源會以下列關鍵字向 Google News 查詢，只抓取符合的報導。用 OR 分隔多個關鍵字，留空代表不篩選。後段仍依覆蓋熱度排名。")
            _kw_global = st.text_area(
                "全球媒體 關鍵字",
                value=_cat_kw.get("全球媒體", DEFAULT_CATEGORY_KEYWORDS.get("全球媒體", "")),
                height=80,
                key="kw_editor_global",
                label_visibility="collapsed",
            )
            if st.button("儲存關鍵字", key="save_kw_global", use_container_width=True):
                _cat_kw["全球媒體"] = _kw_global.strip()
                save_category_keywords(_cat_kw)
                st.success("關鍵字已儲存。")
                st.rerun()
        st.caption(f"共 {len(global_sources_ui)} 筆（唯讀）")
        _global_rows = [source_to_editor_row(x) for x in global_sources_ui]
        _global_df = pd.DataFrame(_global_rows) if _global_rows else pd.DataFrame(columns=["name", "url", "region", "enabled"])
        st.dataframe(_global_df[["name", "url", "region", "enabled"]], use_container_width=True, hide_index=True)

    # ── 中國媒體 ──────────────────────────────────────────────────────────────
    with src_tab_cn:
        with st.expander("🔍 Google News RSS 關鍵字篩選", expanded=False):
            st.caption("此類別的 domain 來源會以下列關鍵字向 Google News 查詢，只抓取符合的報導。用 OR 分隔多個關鍵字，留空代表不篩選。")
            _kw_cn = st.text_area(
                "中國媒體 關鍵字",
                value=_cat_kw.get("中國媒體", DEFAULT_CATEGORY_KEYWORDS.get("中國媒體", "")),
                height=80,
                key="kw_editor_cn",
                label_visibility="collapsed",
            )
            if st.button("儲存關鍵字", key="save_kw_cn", use_container_width=True):
                _cat_kw["中國媒體"] = _kw_cn.strip()
                save_category_keywords(_cat_kw)
                st.success("關鍵字已儲存。")
                st.rerun()
        cn_official_sources = [s for s in fixed_sources if s.get("type") == "cn_official"]
        st.caption(f"共 {len(cn_official_sources)} 筆（唯讀）")
        _cn_rows = [source_to_editor_row(x) for x in cn_official_sources]
        _cn_df = pd.DataFrame(_cn_rows) if _cn_rows else pd.DataFrame(columns=["name", "category", "description"])
        st.dataframe(_cn_df[["name", "category", "description"]], use_container_width=True, hide_index=True)


# =========================================================
# Formats
# =========================================================
elif selected_page == "Formats":

    formats = load_formats()
    if not formats:
        st.info("目前沒有 formats。請先建立 config/formats.json")
    else:
        selected_format = st.selectbox(
            "選擇 Format",
            options=[f.get("name", "default") for f in formats],
            index=0,
            key="formats_selected_name",
        )

        current = next(
            (f for f in formats if f.get("name", "default") == selected_format),
            None
        )

        if current:
            c1, c2 = st.columns(2)

            with c1:
                title_font_size = st.number_input(
                    "Title Font Size",
                    min_value=8,
                    max_value=48,
                    value=int(current.get("title", {}).get("font_size", 16)),
                    step=1,
                )
                title_bold = st.checkbox(
                    "Title Bold",
                    value=bool(current.get("title", {}).get("bold", True)),
                )

                section_font_size = st.number_input(
                    "Section Heading Font Size",
                    min_value=8,
                    max_value=36,
                    value=int(current.get("section_heading", {}).get("font_size", 14)),
                    step=1,
                )
                section_bold = st.checkbox(
                    "Section Heading Bold",
                    value=bool(current.get("section_heading", {}).get("bold", True)),
                )

            with c2:
                body_font_size = st.number_input(
                    "Body Font Size",
                    min_value=8,
                    max_value=24,
                    value=int(current.get("body", {}).get("font_size", 12)),
                    step=1,
                )
                line_spacing = st.number_input(
                    "Line Spacing",
                    min_value=1.0,
                    max_value=3.0,
                    value=float(current.get("body", {}).get("line_spacing", 1.15)),
                    step=0.05,
                )

                notes_style = st.selectbox(
                    "Notes Style",
                    options=["footnote", "endnote", "none"],
                    index=["footnote", "endnote", "none"].index(
                        current.get("notes", {}).get("style", "footnote")
                    ) if current.get("notes", {}).get("style", "footnote") in ["footnote", "endnote", "none"] else 0,
                )

                link_placement = st.selectbox(
                    "Link Placement",
                    options=["inline", "footnote", "none"],
                    index=["inline", "footnote", "none"].index(
                        current.get("links", {}).get("placement", "inline")
                    ) if current.get("links", {}).get("placement", "inline") in ["inline", "footnote", "none"] else 0,
                )

            if st.button("Save Format", key="save_format_btn", use_container_width=True):
                current.setdefault("title", {})
                current.setdefault("section_heading", {})
                current.setdefault("body", {})
                current.setdefault("notes", {})
                current.setdefault("links", {})

                current["title"]["font_size"] = int(title_font_size)
                current["title"]["bold"] = bool(title_bold)

                current["section_heading"]["font_size"] = int(section_font_size)
                current["section_heading"]["bold"] = bool(section_bold)

                current["body"]["font_size"] = int(body_font_size)
                current["body"]["line_spacing"] = float(line_spacing)

                current["notes"]["style"] = notes_style
                current["links"]["placement"] = link_placement

                save_formats(formats)
                st.success("Format 已儲存。")

# =========================================================
# Automation
# =========================================================
# =========================================================
# Automation
# =========================================================
elif selected_page == "Schedule":
    import pandas as pd
    from datetime import datetime
    from utils.auto_export import compute_next_runs


    config = load_auto_export()
    config.setdefault("enabled", True)
    config.setdefault("schedules", [])
    config.setdefault("drive_folders", [])

    schedules = config["schedules"]

    # -------------------------
    # 載入來源 / 專家 / 模板
    # -------------------------
    all_sources = load_sources()
    all_experts = load_experts()
    all_profiles = load_profiles()

    source_categories = get_source_categories(all_sources)
    expert_categories = get_expert_categories(all_experts)

    profile_names = []
    for p in all_profiles:
        if isinstance(p, dict):
            name = (p.get("name") or "").strip()
            if name:
                profile_names.append(name)
    if "default" not in profile_names:
        profile_names = ["default"] + profile_names

    source_name_options = []
    for s in all_sources:
        n = (s.get("name") or "").strip()
        if n and n not in source_name_options:
            source_name_options.append(n)

    expert_name_options = []
    for e in all_experts:
        n = display_expert_name(e)
        if n and n not in expert_name_options:
            expert_name_options.append(n)

    # -------------------------
    # helpers
    # -------------------------
    def _csv_to_list(text):
        if text is None:
            return []
        text = str(text).strip()
        if not text:
            return []
        for sep in ["\n", "；", ";", "、"]:
            text = text.replace(sep, ",")
        return [x.strip() for x in text.split(",") if x.strip()]

    def _safe_int(v, default):
        try:
            return int(v)
        except Exception:
            return default

    def _schedule_to_table_row(s):
        mode = s.get("schedule_mode", "daily")

        if mode == "once":
            time_or_interval = s.get("once_datetime", "")
        elif mode == "hourly":
            time_or_interval = f'{s.get("hourly_interval_hours", 1)}h'
        elif mode == "daily":
            time_or_interval = ", ".join(s.get("daily_times", []))
        elif mode == "weekly":
            days = ",".join([str(x) for x in s.get("weekly_days", [])])
            times = ",".join(s.get("weekly_times", []))
            time_or_interval = f"{days} @ {times}"
        elif mode == "monthly":
            days = ",".join([str(x) for x in s.get("monthly_days", [])])
            times = ",".join(s.get("monthly_times", []))
            time_or_interval = f"{days} @ {times}"
        else:
            time_or_interval = ""

        return {
            "name": s.get("name", ""),
            "schedule_mode": mode,
            "time_or_interval": time_or_interval,
            "language": s.get("language", "繁體中文"),
            "profile": s.get("profile", "default"),
            "format_name": s.get("format_name", "default"),
            "output_formats": ", ".join(s.get("output_formats", ["docx"])),
            "output_targets": ", ".join(s.get("output_targets", ["local"])),
            "coverage_hours": s.get("coverage_hours", 24),
            "delete": False,
        }

    def _table_row_to_schedule(row, base_schedule):
        s = dict(base_schedule)

        s["name"] = str(row.get("name", "")).strip() or "Untitled Schedule"
        s["schedule_mode"] = str(row.get("schedule_mode", "daily")).strip() or "daily"
        s["language"] = str(row.get("language", "繁體中文")).strip() or "繁體中文"
        s["profile"] = str(row.get("profile", "default")).strip() or "default"
        s["format_name"] = str(row.get("format_name", "default")).strip() if "format_name" in row else "default"
        s["output_formats"] = _csv_to_list(row.get("output_formats", "docx")) or ["docx"]
        s["output_targets"] = _csv_to_list(row.get("output_targets", "local")) or ["local"]
        s["coverage_hours"] = max(1, _safe_int(row.get("coverage_hours", 24), 24))

        mode = s["schedule_mode"]
        raw = str(row.get("time_or_interval", "")).strip()

        # 先保留舊值，避免編輯表格時誤清空其他模式欄位
        s.setdefault("once_datetime", "")
        s.setdefault("hourly_interval_hours", 1)
        s.setdefault("daily_times", ["09:00"])
        s.setdefault("weekly_days", [0])
        s.setdefault("weekly_times", ["09:00"])
        s.setdefault("monthly_days", [1])
        s.setdefault("monthly_times", ["09:00"])

        if mode == "once":
            s["once_datetime"] = raw

        elif mode == "hourly":
            txt = raw.lower().replace("hours", "").replace("hour", "").replace("h", "").strip()
            s["hourly_interval_hours"] = max(1, _safe_int(txt, s.get("hourly_interval_hours", 1)))

        elif mode == "daily":
            s["daily_times"] = _csv_to_list(raw) or s.get("daily_times", ["09:00"])

        elif mode == "weekly":
            # 格式: 1,3,5 @ 09:00,18:00
            if "@" in raw:
                left, right = raw.split("@", 1)
                weekly_days = []
                for x in _csv_to_list(left):
                    try:
                        d = int(x)
                        if 0 <= d <= 6:
                            weekly_days.append(d)
                    except Exception:
                        pass
                weekly_times = _csv_to_list(right)
                s["weekly_days"] = weekly_days or s.get("weekly_days", [0])
                s["weekly_times"] = weekly_times or s.get("weekly_times", ["09:00"])

        elif mode == "monthly":
            # 格式: 1,15,28 @ 09:00,18:00
            if "@" in raw:
                left, right = raw.split("@", 1)
                monthly_days = []
                for x in _csv_to_list(left):
                    try:
                        d = int(x)
                        if 1 <= d <= 31:
                            monthly_days.append(d)
                    except Exception:
                        pass
                monthly_times = _csv_to_list(right)
                s["monthly_days"] = monthly_days or s.get("monthly_days", [1])
                s["monthly_times"] = monthly_times or s.get("monthly_times", ["09:00"])

        return s

    def _new_schedule():
        return {
            "name": f"Briefing {len(config['schedules']) + 1}",
            "schedule_mode": "once",
            "once_datetime": now_tw().strftime("%Y-%m-%d %H:%M"),
            "hourly_interval_hours": 4,
            "daily_times": ["09:00"],
            "weekly_days": [0],
            "weekly_times": ["09:00"],
            "monthly_days": [1],
            "monthly_times": ["09:00"],
            "coverage_hours": 24,
            "selected_source_categories": [],
            "selected_source_names": [],
            "selected_expert_categories": [],
            "selected_expert_names": [],
            "language": "繁體中文",
            "profile": "default",
            "format_name": "default",
            "output_formats": ["docx"],
            "output_targets": ["local"],
            "google_drive_folder_id": "",
        }

    if "automation_selected_index" not in st.session_state:
        st.session_state.automation_selected_index = 0

    # 版本計數器：每次新增/刪除排程後加一，強制 data_editor 重置為最新資料，
    # 防止 session_state 舊狀態覆蓋掉剛新增的排程。
    _sch_v = st.session_state.get("_sch_version", 0)

    # -------------------------
    # 啟用開關
    # -------------------------
    top_c1, top_c2, top_c3 = st.columns([1, 1, 1])

    with top_c1:
        config["enabled"] = st.checkbox("啟用自動排程", value=config.get("enabled", True))

    with top_c2:
        if st.button("➕ 新增排程", use_container_width=True):
            config["schedules"].append(_new_schedule())
            save_auto_export(config)
            st.session_state.automation_selected_index = max(0, len(config["schedules"]) - 1)
            st.session_state["_sch_version"] = _sch_v + 1
            st.rerun()

    with top_c3:
        if st.button("💾 儲存全部排程", use_container_width=True):
            save_auto_export(config)
            st.success("已儲存排程設定")

    # -------------------------
    # 排程表
    # -------------------------
    st.markdown("### 排程列表")

    table_rows = [_schedule_to_table_row(s) for s in schedules]
    if not table_rows:
        table_rows = [{
            "name": "",
            "schedule_mode": "once",
            "time_or_interval": now_tw().strftime("%Y-%m-%d %H:%M"),
            "language": "繁體中文",
            "profile": "default",
            "format_name": "default",
            "output_formats": "docx",
            "output_targets": "local",
            "coverage_hours": 24,
            "delete": False,
        }]

    # 加入「下次執行」欄位
    for i, row in enumerate(table_rows):
        if i < len(schedules):
            _next = compute_next_runs(schedules[i], 1, now=now_tw())
            row["next_run"] = _next[0].strftime("%m/%d %H:%M") if _next else "—（已過期）"
        else:
            row["next_run"] = "—"

    table_df = pd.DataFrame(table_rows)

    edited_df = st.data_editor(
        table_df,
        use_container_width=True,
        num_rows="fixed",
        column_config={
            "name": st.column_config.TextColumn("排程名稱"),
            "schedule_mode": st.column_config.SelectboxColumn(
                "模式",
                options=["once", "hourly", "daily", "weekly", "monthly"],
            ),
            "time_or_interval": st.column_config.TextColumn("時間 / 間隔"),
            "language": st.column_config.SelectboxColumn(
                "語言",
                options=["繁體中文", "英文", "日文", "簡體中文"],
            ),
            "profile": st.column_config.SelectboxColumn(
                "Profile",
                options=profile_names,
            ),
            "format_name": st.column_config.SelectboxColumn(
                "Format",
                options=format_names,
            ),
            "output_formats": st.column_config.TextColumn("格式"),
            "output_targets": st.column_config.TextColumn("輸出位置"),
            "coverage_hours": st.column_config.NumberColumn("涵蓋小時", min_value=1, step=1),
            "next_run": st.column_config.TextColumn("下次執行", disabled=True),
            "delete": st.column_config.CheckboxColumn("刪除"),
        },
        key=f"automation_schedule_editor_{_sch_v}",
    )

    if len(schedules) > 0:
        rebuilt = []
        for i, row in edited_df.iterrows():
            if bool(row.get("delete", False)):
                continue
            base_schedule = schedules[i] if i < len(schedules) else _new_schedule()
            rebuilt.append(_table_row_to_schedule(dict(row), base_schedule))

        # 安全網：若 data_editor 的 session state 是舊的（列數少於 schedules），
        # 保留那些未顯示在 editor 裡的排程，避免它們被靜默覆蓋。
        editor_row_count = len(edited_df)
        if editor_row_count < len(schedules):
            for extra_i in range(editor_row_count, len(schedules)):
                rebuilt.append(schedules[extra_i])

        config["schedules"] = rebuilt
        schedules = config["schedules"]

        if st.session_state.automation_selected_index >= len(schedules):
            st.session_state.automation_selected_index = max(0, len(schedules) - 1)

    # -------------------------
    # 細部設定
    # -------------------------
    st.markdown("### 排程細部設定")

    if not schedules:
        st.info("目前沒有排程。請先按「新增排程」。")
    else:
        selected_options = [
            f"{idx + 1}. {s.get('name', 'Untitled Schedule')}"
            for idx, s in enumerate(schedules)
        ]

        selected_label = st.selectbox(
            "選擇排程",
            options=selected_options,
            index=min(st.session_state.automation_selected_index, len(selected_options) - 1),
        )
        selected_idx = selected_options.index(selected_label)
        st.session_state.automation_selected_index = selected_idx

        s = schedules[selected_idx]

        c1, c2 = st.columns(2)

        with c1:
            s["name"] = st.text_input("排程名稱", value=s.get("name", ""), key=f"name_{selected_idx}")
            s["coverage_hours"] = int(
                st.number_input(
                    "涵蓋過去幾小時",
                    min_value=1,
                    max_value=720,
                    value=int(s.get("coverage_hours", 24)),
                    step=1,
                    key=f"coverage_{selected_idx}",
                )
            )
            s["language"] = st.selectbox(
                "語言",
                options=["繁體中文", "英文", "日文", "簡體中文"],
                index=["繁體中文", "英文", "日文", "簡體中文"].index(
                    s.get("language", "繁體中文")
                ) if s.get("language", "繁體中文") in ["繁體中文", "英文", "日文", "簡體中文"] else 0,
                key=f"language_{selected_idx}",
            )
            s["profile"] = st.selectbox(
                "模板",
                options=profile_names,
                index=profile_names.index(s.get("profile", "default"))
                if s.get("profile", "default") in profile_names else 0,
                key=f"profile_{selected_idx}",
            )

            s["format_name"] = st.selectbox(
                "Format",
                options=format_names,
                index=format_names.index(s.get("format_name", "default"))
                if s.get("format_name", "default") in format_names else 0,
                key=f"format_{selected_idx}",
            )
            s["output_formats"] = st.multiselect(
                "輸出格式",
                options=["docx", "pdf"],
                default=s.get("output_formats", ["docx"]),
                key=f"formats_{selected_idx}",
            )
            s["output_targets"] = st.multiselect(
                "輸出位置",
                options=["local", "google_drive"],
                default=s.get("output_targets", ["local"]),
                key=f"targets_{selected_idx}",
            )
            # Google Drive 資料夾選擇
            _df_list = config.get("drive_folders", [])
            _df_names = [f.get("name", "") or f.get("folder_id", "") for f in _df_list]
            _cur_fid = s.get("google_drive_folder_id", "")
            # 找目前 folder_id 對應的名稱
            _cur_idx = next(
                (i for i, f in enumerate(_df_list) if f.get("folder_id") == _cur_fid),
                None
            )
            if _df_names:
                _options = ["（手動輸入）"] + _df_names
                _sel = st.selectbox(
                    "Google Drive 資料夾",
                    options=_options,
                    index=(_cur_idx + 1) if _cur_idx is not None else 0,
                    key=f"gdrive_sel_{selected_idx}",
                )
                if _sel == "（手動輸入）":
                    s["google_drive_folder_id"] = st.text_input(
                        "Folder ID（手動輸入）",
                        value=_cur_fid,
                        key=f"gdrive_manual_{selected_idx}",
                    )
                else:
                    _chosen = next((f for f in _df_list if (f.get("name") or f.get("folder_id")) == _sel), None)
                    s["google_drive_folder_id"] = _chosen.get("folder_id", "") if _chosen else ""
                    st.caption(f"Folder ID：`{s['google_drive_folder_id']}`")
            else:
                s["google_drive_folder_id"] = st.text_input(
                    "Google Drive Folder ID",
                    value=_cur_fid,
                    key=f"gdrive_{selected_idx}",
                    help="先在上方「Google Drive 資料夾」區塊新增資料夾，之後可在此選擇。",
                )

        with c2:
            s["schedule_mode"] = st.selectbox(
                "排程模式",
                options=["once", "hourly", "daily", "weekly", "monthly"],
                index=["once", "hourly", "daily", "weekly", "monthly"].index(
                    s.get("schedule_mode", "daily")
                ),
                key=f"mode_{selected_idx}",
            )

            mode = s["schedule_mode"]

            if mode == "once":
                s["once_datetime"] = st.text_input(
                    "指定時間（YYYY-MM-DD HH:MM）",
                    value=s.get("once_datetime") or now_tw().strftime("%Y-%m-%d %H:%M"),
                    key=f"once_{selected_idx}",
                )

            elif mode == "hourly":
                s["hourly_interval_hours"] = int(
                    st.number_input(
                        "每幾小時執行一次",
                        min_value=1,
                        max_value=168,
                        value=int(s.get("hourly_interval_hours", 4)),
                        step=1,
                        key=f"hourly_{selected_idx}",
                    )
                )
                _sf_raw = s.get("start_from", "")
                try:
                    _sf_dt = datetime.strptime(_sf_raw[:16], "%Y-%m-%d %H:%M") if len(_sf_raw) >= 16 else (datetime.strptime(_sf_raw[:10], "%Y-%m-%d") if _sf_raw else now_tw())
                except Exception:
                    _sf_dt = now_tw()
                _sfc1, _sfc2 = st.columns(2)
                with _sfc1:
                    _sf_date = st.date_input("生效開始日期", value=_sf_dt.date(), key=f"start_from_date_{selected_idx}")
                with _sfc2:
                    _sf_time = st.time_input("生效開始時間", value=_sf_dt.time().replace(second=0, microsecond=0), key=f"start_from_time_{selected_idx}", step=300)
                s["start_from"] = f"{_sf_date.isoformat()} {_sf_time.strftime('%H:%M')}"

            elif mode == "daily":
                daily_text = st.text_input(
                    "每日時間（可多個，以逗號分隔）",
                    value=", ".join(s.get("daily_times", ["09:00"])),
                    key=f"daily_{selected_idx}",
                )
                s["daily_times"] = _csv_to_list(daily_text) or ["09:00"]
                _sf_raw = s.get("start_from", "")
                try:
                    _sf_dt = datetime.strptime(_sf_raw[:16], "%Y-%m-%d %H:%M") if len(_sf_raw) >= 16 else (datetime.strptime(_sf_raw[:10], "%Y-%m-%d") if _sf_raw else now_tw())
                except Exception:
                    _sf_dt = now_tw()
                _sfc1, _sfc2 = st.columns(2)
                with _sfc1:
                    _sf_date = st.date_input("生效開始日期", value=_sf_dt.date(), key=f"start_from_date_{selected_idx}")
                with _sfc2:
                    _sf_time = st.time_input("生效開始時間", value=_sf_dt.time().replace(second=0, microsecond=0), key=f"start_from_time_{selected_idx}", step=300)
                s["start_from"] = f"{_sf_date.isoformat()} {_sf_time.strftime('%H:%M')}"

            elif mode == "weekly":
                weekly_days = st.multiselect(
                    "每週星期幾（0=Mon, 6=Sun）",
                    options=[0, 1, 2, 3, 4, 5, 6],
                    default=s.get("weekly_days", [0]),
                    key=f"weekly_days_{selected_idx}",
                )
                weekly_times_text = st.text_input(
                    "每週時間（可多個，以逗號分隔）",
                    value=", ".join(s.get("weekly_times", ["09:00"])),
                    key=f"weekly_times_{selected_idx}",
                )
                s["weekly_days"] = weekly_days or [0]
                s["weekly_times"] = _csv_to_list(weekly_times_text) or ["09:00"]
                _sf_raw = s.get("start_from", "")
                try:
                    _sf_dt = datetime.strptime(_sf_raw[:16], "%Y-%m-%d %H:%M") if len(_sf_raw) >= 16 else (datetime.strptime(_sf_raw[:10], "%Y-%m-%d") if _sf_raw else now_tw())
                except Exception:
                    _sf_dt = now_tw()
                _sfc1, _sfc2 = st.columns(2)
                with _sfc1:
                    _sf_date = st.date_input("生效開始日期", value=_sf_dt.date(), key=f"start_from_date_{selected_idx}")
                with _sfc2:
                    _sf_time = st.time_input("生效開始時間", value=_sf_dt.time().replace(second=0, microsecond=0), key=f"start_from_time_{selected_idx}", step=300)
                s["start_from"] = f"{_sf_date.isoformat()} {_sf_time.strftime('%H:%M')}"

            elif mode == "monthly":
                monthly_days_text = st.text_input(
                    "每月幾號（可多個，以逗號分隔）",
                    value=", ".join([str(x) for x in s.get("monthly_days", [1])]),
                    key=f"monthly_days_{selected_idx}",
                )
                monthly_times_text = st.text_input(
                    "每月時間（可多個，以逗號分隔）",
                    value=", ".join(s.get("monthly_times", ["09:00"])),
                    key=f"monthly_times_{selected_idx}",
                )
                month_days = []
                for x in _csv_to_list(monthly_days_text):
                    try:
                        d = int(x)
                        if 1 <= d <= 31:
                            month_days.append(d)
                    except Exception:
                        pass
                s["monthly_days"] = month_days or [1]
                s["monthly_times"] = _csv_to_list(monthly_times_text) or ["09:00"]
                _sf_raw = s.get("start_from", "")
                try:
                    _sf_dt = datetime.strptime(_sf_raw[:16], "%Y-%m-%d %H:%M") if len(_sf_raw) >= 16 else (datetime.strptime(_sf_raw[:10], "%Y-%m-%d") if _sf_raw else now_tw())
                except Exception:
                    _sf_dt = now_tw()
                _sfc1, _sfc2 = st.columns(2)
                with _sfc1:
                    _sf_date = st.date_input("生效開始日期", value=_sf_dt.date(), key=f"start_from_date_{selected_idx}")
                with _sfc2:
                    _sf_time = st.time_input("生效開始時間", value=_sf_dt.time().replace(second=0, microsecond=0), key=f"start_from_time_{selected_idx}", step=300)
                s["start_from"] = f"{_sf_date.isoformat()} {_sf_time.strftime('%H:%M')}"

        st.markdown("#### 報告模式")
        def _sched_rmode_label(x):
            if x == "single":
                return "單份報告"
            elif x == "multi_phase":
                return "綜合報告（多段生成）"
            else:
                return "分段報告（按章節搜尋）"
        _rmode_options = ["single", "multi_phase", "segmented"]
        _rmode_current = s.get("report_mode", "single")
        if _rmode_current not in _rmode_options:
            _rmode_current = "single"
        s["report_mode"] = st.radio(
            "報告模式",
            options=_rmode_options,
            format_func=_sched_rmode_label,
            index=_rmode_options.index(_rmode_current),
            horizontal=True,
            key=f"report_mode_{selected_idx}",
        )
        if s["report_mode"] == "multi_phase":
            _gmap_s = {k: report_engine._MULTIPHASE_GROUP_ZH.get(k, k)
                       for k in report_engine.MULTIPHASE_GROUP_OPTIONS}
            s["multiphase_groups"] = st.multiselect(
                "包含來源群組（空白 = 全部）",
                options=list(_gmap_s.keys()),
                format_func=lambda x: _gmap_s.get(x, x),
                default=s.get("multiphase_groups") or [],
                key=f"multiphase_groups_{selected_idx}",
            )
        elif s["report_mode"] == "segmented":
            s["multiphase_groups"] = []
            st.caption("分段報告：每章每節獨立搜尋 Google News，無需選擇來源群組。")
        else:
            s["multiphase_groups"] = []

        st.markdown("#### 來源與專家篩選")

        c3, c4 = st.columns(2)

        with c3:
            s["selected_source_categories"] = st.multiselect(
                "來源分類",
                options=source_categories,
                default=s.get("selected_source_categories", []),
                key=f"src_cat_{selected_idx}",
            )

            filtered_source_name_options = []
            for item in all_sources:
                item_name = (item.get("name") or "").strip()
                item_cats = item.get("category", []) or []
                if not s["selected_source_categories"]:
                    if item_name and item_name not in filtered_source_name_options:
                        filtered_source_name_options.append(item_name)
                else:
                    if any(cat in s["selected_source_categories"] for cat in item_cats):
                        if item_name and item_name not in filtered_source_name_options:
                            filtered_source_name_options.append(item_name)

            existing_source_names = [
                x for x in s.get("selected_source_names", [])
                if x in filtered_source_name_options
            ]

            s["selected_source_names"] = st.multiselect(
                "來源名稱",
                options=filtered_source_name_options,
                default=existing_source_names,
                key=f"src_names_{selected_idx}",
            )

        with c4:
            s["selected_expert_categories"] = st.multiselect(
                "專家分類",
                options=expert_categories,
                default=s.get("selected_expert_categories", []),
                key=f"exp_cat_{selected_idx}",
            )

            filtered_expert_name_options = []
            for item in all_experts:
                item_name = (item.get("name") or "").strip()
                item_cats = item.get("category", []) or []
                if not s["selected_expert_categories"]:
                    if item_name and item_name not in filtered_expert_name_options:
                        filtered_expert_name_options.append(item_name)
                else:
                    if any(cat in s["selected_expert_categories"] for cat in item_cats):
                        if item_name and item_name not in filtered_expert_name_options:
                            filtered_expert_name_options.append(item_name)

            existing_expert_names = [
                x for x in s.get("selected_expert_names", [])
                if x in filtered_expert_name_options
            ]

            s["selected_expert_names"] = st.multiselect(
                "專家名稱",
                options=filtered_expert_name_options,
                default=existing_expert_names,
                key=f"exp_names_{selected_idx}",
            )

        config["schedules"][selected_idx] = s

        save_c1, save_c2 = st.columns(2)

        with save_c1:
            if st.button("儲存此排程", key=f"save_this_{selected_idx}", use_container_width=True):
                save_auto_export(config)
                st.success("已儲存此排程")

        with save_c2:
            if st.button("刪除此排程", key=f"delete_this_{selected_idx}", use_container_width=True):
                config["schedules"].pop(selected_idx)
                save_auto_export(config)
                st.session_state.automation_selected_index = max(0, selected_idx - 1)
                st.rerun()

    # -------------------------
    # 執行狀態與歷史
    # -------------------------
    st.markdown("### 執行狀態")

    _exec_state = load_auto_export_state()
    _running_now = _exec_state.get("running_now", [])
    _running_started = _exec_state.get("running_started_at", {})

    if _running_now:
        for _rn in _running_now:
            _started = _running_started.get(_rn, "")
            st.warning(f"⏳ **{_rn}** 正在執行中…（開始時間：{_started}）")
    else:
        st.success("✅ 目前沒有排程正在執行")

    st.markdown("### 執行歷史")

    _run_history = _exec_state.get("run_history", [])
    if not _run_history:
        st.info("尚無執行紀錄。")
    else:
        if st.button("🔄 重新整理", key="refresh_history"):
            st.rerun()
        for _h in _run_history:
            _h_ok = _h.get("ok", False)
            _h_name = _h.get("name", "")
            _h_started = _h.get("started_at", "")
            _h_duration = _h.get("duration_sec", 0)
            _h_msg = _h.get("message", "")
            _icon = "✅" if _h_ok else "❌"
            with st.expander(f"{_icon} {_h_name}　{_h_started}　（{_h_duration} 秒）", expanded=False):
                if _h_msg:
                    st.caption(_h_msg)

    # -------------------------
    # Google Drive 資料夾管理
    # -------------------------
    st.divider()
    st.markdown("### Google Drive 資料夾")

    drive_folders = config.get("drive_folders", [])

    # 顯示現有資料夾列表
    if drive_folders:
        for _dfi, _df in enumerate(drive_folders):
            _dfc1, _dfc2, _dfc3 = st.columns([3, 5, 1])
            with _dfc1:
                _new_name = st.text_input(
                    "資料夾名稱",
                    value=_df.get("name", ""),
                    key=f"df_name_{_dfi}",
                    label_visibility="collapsed",
                    placeholder="資料夾名稱",
                )
            with _dfc2:
                _new_fid = st.text_input(
                    "Folder ID",
                    value=_df.get("folder_id", ""),
                    key=f"df_fid_{_dfi}",
                    label_visibility="collapsed",
                    placeholder="Google Drive Folder ID",
                )
            with _dfc3:
                if st.button("🗑️", key=f"df_del_{_dfi}", use_container_width=True):
                    config["drive_folders"].pop(_dfi)
                    save_auto_export(config)
                    st.rerun()
            # 若有修改則即時更新
            if _new_name != _df.get("name", "") or _new_fid != _df.get("folder_id", ""):
                config["drive_folders"][_dfi] = {"name": _new_name, "folder_id": _new_fid}
    else:
        st.caption("尚未設定任何資料夾。")

    if st.button("＋ 新增資料夾", key="add_drive_folder"):
        config["drive_folders"].append({"name": "", "folder_id": ""})
        save_auto_export(config)
        st.rerun()

    if drive_folders and st.button("💾 儲存資料夾設定", key="save_drive_folders"):
        save_auto_export(config)
        st.success("已儲存資料夾設定")


# =========================================================
# Reports
# =========================================================
elif selected_page == "Reports":

    report_files = sorted(OUTPUT_DIR.glob("*"), key=lambda f: f.stat().st_mtime, reverse=True)
    report_files = [f for f in report_files if f.is_file()]

    if not report_files:
        st.info("目前沒有已儲存的報告。生成報告後會自動出現在這裡。")
    else:
        st.caption(f"共 {len(report_files)} 個檔案，儲存於 outputs 資料夾")
        st.markdown("---")

        _reports_drive_folder_id = auto_export_cfg.get("default_drive_folder_id", "")

        for _rf in report_files:
            _stat = _rf.stat()
            _size_kb = _stat.st_size / 1024
            _mtime = datetime.fromtimestamp(_stat.st_mtime).strftime("%Y-%m-%d %H:%M")

            col_name, col_dl, col_drive, col_del = st.columns([5, 1, 1, 1])

            with col_name:
                st.write(f"**{_rf.name}**")
                st.caption(f"{_mtime} · {_size_kb:.1f} KB")

            with col_dl:
                with open(_rf, "rb") as _fp:
                    st.download_button(
                        "⬇ 下載",
                        data=_fp.read(),
                        file_name=_rf.name,
                        key=f"dl_{_rf.name}",
                        use_container_width=True,
                    )

            with col_drive:
                if st.button("☁ Drive", key=f"drive_{_rf.name}", use_container_width=True):
                    _uploaded, _err = _try_upload_to_drive(str(_rf), _reports_drive_folder_id)
                    if _err:
                        st.error(f"上傳失敗：{_err}")
                    else:
                        st.success(f"已上傳！")

            with col_del:
                if st.button("🗑 刪除", key=f"del_{_rf.name}", use_container_width=True):
                    _rf.unlink()
                    st.rerun()

            st.markdown("---")